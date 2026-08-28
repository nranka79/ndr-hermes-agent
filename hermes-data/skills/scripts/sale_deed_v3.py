from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING

doc = Document()

# Page setup — A4, narrow margins
section = doc.sections[0]
section.page_width = Inches(8.27)
section.page_height = Inches(11.69)
section.left_margin = Inches(1.0)
section.right_margin = Inches(1.0)
section.top_margin = Inches(0.8)
section.bottom_margin = Inches(0.8)

# Default style
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(11)

# ─── Helpers ──────────────────────────────────────────

def sp(doc, text="", indent=0, bold=False, bold_phrases=None,
       space_b=2, space_a=2, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    p = doc.add_paragraph()
    p.alignment = align
    pf = p.paragraph_format
    pf.space_before = Pt(space_b)
    pf.space_after = Pt(space_a)
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    if indent:
        pf.left_indent = Inches(indent)
    if bold:
        r = p.add_run(text); r.bold = True; r.font.size = Pt(11)
    elif bold_phrases and text:
        remaining = text
        for bp in bold_phrases:
            idx = remaining.find(bp)
            if idx >= 0:
                if idx > 0:
                    r = p.add_run(remaining[:idx]); r.font.size = Pt(11)
                r2 = p.add_run(bp); r2.bold = True; r2.font.size = Pt(11)
                remaining = remaining[idx+len(bp):]
        if remaining:
            r = p.add_run(remaining); r.font.size = Pt(11)
    else:
        if text:
            r = p.add_run(text); r.font.size = Pt(11)
    return p

def heading(doc, text, size=12, bold=True, space_b=10, space_a=4,
            align=WD_ALIGN_PARAGRAPH.LEFT, underline=False):
    p = doc.add_paragraph()
    p.alignment = align
    pf = p.paragraph_format
    pf.space_before = Pt(space_b)
    pf.space_after = Pt(space_a)
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    r = p.add_run(text); r.bold = bold; r.font.size = Pt(size)
    if underline:
        r.underline = True
    return p

def sub_para(doc, text, indent=0.5, space_b=3, space_a=3):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = p.paragraph_format
    pf.left_indent = Inches(indent)
    pf.first_line_indent = Inches(-0.3)
    pf.space_before = Pt(space_b)
    pf.space_after = Pt(space_a)
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    r = p.add_run(text); r.font.size = Pt(11)
    return p

def divider(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.space_before = Pt(3)
    pf.space_after = Pt(3)
    r = p.add_run("- - -"); r.font.size = Pt(10)

def kv(doc, key, value, indent=0.6):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = p.paragraph_format
    pf.left_indent = Inches(indent)
    pf.space_before = Pt(1)
    pf.space_after = Pt(1)
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    r1 = p.add_run(key + " "); r1.bold = True; r1.font.size = Pt(11)
    r2 = p.add_run(value); r2.font.size = Pt(11)

# ─── TITLE ──────────────────────────────────────────────

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(0)
p.paragraph_format.space_after = Pt(2)
r = p.add_run("ABSOLUTE SALE DEED"); r.bold = True; r.font.size = Pt(16)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(0)
p.paragraph_format.space_after = Pt(2)
r = p.add_run("India - TN RERA Compliant Draft"); r.font.size = Pt(11)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(0)
p.paragraph_format.space_after = Pt(4)
r = p.add_run("Document Writer Attested"); r.italic = True; r.font.size = Pt(10)

divider(doc)

sp(doc, 'THIS DEED OF ABSOLUTE SALE is made and executed at Bagalur, Krishnagiri District, Tamil Nadu, ON THIS THE _____ DAY OF MAY 2026.',
   bold_phrases=["ABSOLUTE SALE"], space_b=6, space_a=4)

# ─── PARTIES ─────────────────────────────────────────

heading(doc, "BETWEEN", size=12, space_b=8, space_a=4)

sp(doc, 'M/s. DRA THINDLU LAND PARTNERS, a Partnership Firm registered under the Indian Partnership Act, 1932, having its registered office at Queens Corner, 3rd Floor, No.302A, Queens Road, Bengaluru East, Karnataka - 560001 (PAN: AAXFD2296G; Firm Registration No.: SJN-F655-2024-25).',
   bold_phrases=["M/s. DRA THINDLU LAND PARTNERS"], space_b=4, space_a=3)

sp(doc, 'Hereafter referred to as the "VENDOR", which expression shall, unless it be repugnant to the context or meaning thereof, include its partners, successors, heirs, legal representatives, administrators, executors, and permitted assigns.',
   space_b=2, space_a=4)

sp(doc, 'The VENDOR is represented by its current partners, as follows:', bold_phrases=["VENDOR"], space_b=2, space_a=3)

sub_para(doc, "(i)  Partner 1 - Continuing Partner:  M/s. DRA REALTY PRIVATE LIMITED, a company incorporated under the Companies Act, 2013 (CIN: U70100KA2011PTC058105), having its registered office at No.201A/202BA, Queens Corner, 3 Queens Road, Bangalore - 560001, represented by its Director, Mr. KISHAN MURJANI NAIR, S/o Mr. Pukhraj Murjani Nair, aged about 31 years, holding Aadhaar No. 970838057634, and residing at No.302A, Queens Corner, 3 Queens Road, Bengaluru - 560001.")

sub_para(doc, "(ii)  Partner 2 - Incoming Partner:  Mr. NISHANT RANKA, S/o Late Sri Dinesh Ranka, aged about 46 years, holding Aadhaar No. 4159 0535 2796, and residing at No.302A, Queens Corner, 3 Queens Road, Bengaluru - 560001.")

sp(doc, 'Hereafter collectively referred to as the "Partners of the Vendor".',
   bold_phrases=["Partners of the Vendor"], space_b=4, space_a=4)

sp(doc, 'MR. MANJUNATH SINGH MANOHAR SINGH, S/o Mr. Manjunath Singh, aged about 53 years, born on 24th February 1973 at Bangalore, Karnataka, residing at Villa No. H, MIMS Espaccio, Survey No.105/8 & 106/1, Behind Jakkur Flying Club, Jakkur, Bangalore North - 560064, Karnataka, India, holding PAN: AOQPS1456J, Aadhaar No: 4480 4783 5557, Passport No: Z2448430.',
   bold_phrases=["MR. MANJUNATH SINGH MANOHAR SINGH"], space_b=4, space_a=3)

sp(doc, 'Hereafter referred to as the "VENDEE", which expression shall, unless it be repugnant to the context or meaning thereof, include his heirs, executors, administrators, legal representatives, and assigns.',
   space_b=2, space_a=3)

sp(doc, 'The VENDOR and the VENDEE are hereafter collectively referred to as the "Parties" and individually as a "Party".',
   bold_phrases=["Parties", "Party"], space_b=2, space_a=6)

divider(doc)

# ─── PART I ───────────────────────────────────────────

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(12)
p.paragraph_format.space_after = Pt(6)
r = p.add_run("PART I -- BACKGROUND AND RECITALS"); r.bold = True; r.font.size = Pt(13)

# A
heading(doc, "A.  PARTNERSHIP RECONSTITUTION", size=12, space_b=8, space_a=4)

sp(doc, 'The Vendor, M/s. DRA THINDLU LAND PARTNERS (hereafter referred to as "the Firm" or "the Vendor"), was originally constituted under a Partnership Deed dated 26th September 2024, between M/s. DRA Realty Pvt. Ltd. (represented by Director Mr. Nishant Ranka) as Managing Partner, and Mr. Srinivas Krishnappa (Aadhaar: 4487 9498 6704) as Profit Share Partner, for the purpose of real estate business including acquisition, development, and sale of land parcels and joint development agreements.',
   bold_phrases=["DRA THINDLU LAND PARTNERS"], space_b=2, space_a=4)

sp(doc, 'The said Partnership was subsequently reconstituted by a Deed of Reconstitution of Partnership dated 1st August 2025, executed on an e-Stamp Certificate bearing No. IN-KA06881622484221X (stamp duty of Rs.2,000/- paid by Nishant Ranka), and filed with the Registrar of Firms. The reconstitution provides as follows:',
   bold_phrases=["Deed of Reconstitution of Partnership", "1st August 2025", "IN-KA06881622484221X"],
   space_b=2, space_a=3)

sub_para(doc, "(a)  Mr. Srinivas Krishnappa retired as Profit Share Partner and continues under a Separate Personal Arrangement (SPI);")
sub_para(doc, "(b)  M/s. DRA Realty Pvt. Ltd. continued as the First/Continuing Partner; and")
sub_para(doc, "(c)  Mr. Nishant Ranka (Aadhaar: 4159 0535 2796) was admitted as the Second/Incoming Partner.")

sp(doc, 'The Vendor holds a Special Power of Attorney (SPA) executed on 8th May 2026. All future documents executed by the Vendor shall be as per the SPA.',
   bold_phrases=["Special Power of Attorney", "8th May 2026"],
   space_b=4, space_a=4)

# B
heading(doc, "B.  CHAIN OF TITLE", size=12, space_b=8, space_a=4)

sp(doc, 'The schedule property originally formed part of ancestral agricultural land bearing Survey No.240/3, in Sevaganapalli Village, Hosur Taluk, Krishnagiri District, Tamil Nadu, belonging to one Chowda Reddy. Upon the death of Chowda Reddy, the said land devolved upon his legal heirs, including one Subba Reddy (his son), by way of oral family partition -- Subba Reddy retaining the entire Survey No.240/3.',
   space_b=2, space_a=3)

sp(doc, 'Upon the death of Subba Reddy on 03.07.2017, a Legal Heirs Certificate dated 10.06.2024 was issued by the Tahsildar, Hosur, identifying the following legal heirs: Mrs. Nanjamma (wife); Prakash Reddy, Ramesha, Amaresha (sons); C.R. Vinodha (daughter); Venkatasamy Reddy (son-in-law, husband of late Amaravathi); Sheela, V. Sudha, Madhusudan (grandchildren).',
   bold_phrases=["Legal Heirs Certificate"], space_b=2, space_a=3)

sub_para(doc, "(i)  Gift Settlement Deed dated 07.06.2022 -- Doc No.10658/2022, Sub Registrar Hosur: Mrs. Nanjamma (W/o Late Subba Reddy) settled the property by way of gift upon her sons Prakash Reddy, Ramesha, and Amaresha -- no consideration; made out of natural love and affection. Market Value: Rs.9,80,880/-. Boundaries: East: Sy.No.240/4A, 240/4B, 240/4C | West: Sy.No.240/2 | North: Panchayat Road | South: Sy.No.246.")

sub_para(doc, "(ii)  Rectification Deed dated 14.06.2022 -- Doc No.11162/2022, Sub Registrar Hosur: Corrected boundary errors in the Gift Settlement Deed -- North wrongly typed as 'Survey No.226' corrected to 'Panchayat Road'; South wrongly typed as 'Panchayat Road' corrected to 'Survey No.246'.")

sub_para(doc, "(iii)  Registered Partition Deed dated 18.06.2024 -- Doc No.11721/2024, Sub Registrar Hosur: Among the co-owners (Nanjamma, Prakash Reddy, Ramesha, Amaresha, C.R. Vinodha, Venkatasamy Reddy, Sheela, V. Sudha, Madhusudan), the land was divided -- C.R. Vinodha received Ac.0.08 cents (Schedule E); the remaining Ac.1.75 cents in Survey No.240/3 was divided in 1/3rd shares to Prakash Reddy, Ramesha, and Amaresha (Schedules B, C, D).")

sub_para(doc, "(iv)  Survey Sub-division: Survey No.240/3 sub-divided into Survey No.240/3A (Extent: Ac.1.75 cents / Hec.0.70.76). Entry made under Joint Patta No.1665 in the names of Prakash Reddy, Ramesha, and Amaresha.")

sub_para(doc, "(v)  General Power of Attorney dated 29.07.2024 -- Doc No.14749/2024, Sub Registrar Hosur: Mr. Naveen Kumar (S/o Ramesha) executed a GPA in favour of his father Mr. Ramesha, authorising him to sell Ramesha's share on Naveen Kumar's behalf.")

sub_para(doc, "(vi)  Absolute Sale Deed dated 24.10.2024 -- Doc No.20527/2024, Sub Registrar Hosur: The vendors -- Prakash Reddy, Chowda Reddy, Gajendra, Ramesha (with GPA from Naveen Kumar), Arun Kumar, Amaresha, and Harika -- absolutely sold Ac.1.75 cents in Survey No.240/3A to the VENDOR (M/s. DRA Thindlu Land Partners, represented by its partners M/s. DRA Realty Pvt. Ltd. and Mr. Srinivas) for Rs.3,00,00,000/- (Rupees Three Crores Only). Consideration payment receipts dated October 2024, aggregating to Rs.3,00,00,000/-, were executed prior to the execution of this sale deed.")

sp(doc, 'The VENDOR is thus the absolute and marketable owner of Ac.1.75 cents (Hec.0.70.76) in Survey No.240/3A, Sevaganapalli Village, Hosur Taluk, Krishnagiri District, Tamil Nadu, with clear, valid, and marketable title traceable through the chain of documents set out above.',
   bold_phrases=["absolute and marketable owner", "Survey No.240/3A"],
   space_b=4, space_a=4)

# C
heading(doc, "C.  LAYOUT DEVELOPMENT, APPROVALS, AND RERA COMPLIANCE", size=12, space_b=8, space_a=4)

sp(doc, 'The VENDOR developed the said land into a residential plot layout comprising 38 plots, on Survey No.240/3A, Sevaganapalli Village, Hosur Taluk, Krishnagiri District, Tamil Nadu.',
   bold_phrases=["38 plots"], space_b=2, space_a=3)

sp(doc, 'The layout received Technical Approval from the Hosur New Town Development Authority (HNTDA) -- Approval Nos.: SWP Vz: 90/2025 and 38/2025 (Application No.: J5SROIYP/2024). Total layout area: 0.70.76 hectares (7,076.00 sq.m); Net plot area (excluding roads): 4,529.14 sq.m; Total plots: 38. 10% park area of 452.91 sq.m has been paid for at Rs.95,046/- as per guideline value.',
   bold_phrases=["Hosur New Town Development Authority", "SWP Vz: 90/2025 and 38/2025"],
   space_b=2, space_a=3)

sp(doc, 'The VENDOR has gifted the road areas and common areas within the layout to government authorities by registered Gift Deeds:',
   bold_phrases=["Gift Deeds"], space_b=2, space_a=3)

sub_para(doc, "(a)  Gift Deed dated 24.02.2025 -- Doc No.1634/2025, Sub Registrar Bagalur: Roads and common areas gifted to Tamil Nadu Generation and Distribution Corporation Limited (TANGEDCO) -- represented by Superintendent Engineer, Krishnagiri District.")

sub_para(doc, "(b)  Gift Deed dated 24.02.2025 -- Doc No.1632/2025, Sub Registrar Bagalur: Roads and common areas gifted to the Governor of Tamil Nadu -- represented by Panchayat President, Sevaganapalli Village, Sevaganapalli Panchayat, Hosur Union Council.")

sp(doc, 'The layout and the schedule property are situated within the jurisdiction of the Sevaganapalli Panchayat and the Union Council of Hosur.',
   space_b=3, space_a=3)



# D
heading(doc, "D.  AGREEMENT TO SELL AND CONSIDERATION", size=12, space_b=8, space_a=4)

sp(doc, 'The VENDOR had agreed to sell, and the VENDEE had agreed to purchase, the schedule property more particularly described in the Schedule of Property hereunder, for a total sale consideration of Rs.24,00,000/- (Rupees Twenty Four Lakhs Only).',
   bold_phrases=["Rs.24,00,000/-"], space_b=2, space_a=3)

sp(doc, 'The VENDEE has paid, and the VENDOR has received, the entire sale consideration of Rs.24,00,000/- prior to or simultaneously with the execution of this Deed. The VENDOR acknowledges receipt of the same in full, and confirms that no amount is outstanding from the VENDEE towards the purchase of the schedule property.',
   bold_phrases=["acknowledges receipt"], space_b=2, space_a=3)

sp(doc, 'The VENDOR has agreed to sell absolutely to the VENDEE the under-mentioned schedule property for valid and proper sale consideration, and the VENDEE has agreed to purchase the same. In pursuance of the said agreement, the VENDOR and the VENDEE entered into negotiations and arrived at a concluded contract for the sale of the schedule property.',
   space_b=2, space_a=4)

divider(doc)

# ─── PART II ───────────────────────────────────────────

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(12)
p.paragraph_format.space_after = Pt(6)
r = p.add_run("PART II -- OPERATIVE PROVISIONS"); r.bold = True; r.font.size = Pt(13)

# Clause 1
heading(doc, "CLAUSE 1 -- TRANSFER OF TITLE (ABSOLUTE SALE)", size=12, space_b=8, space_a=4)

sp(doc, 'In pursuance of the said agreement and in consideration of the sum of Rs.24,00,000/- (Rupees Twenty Four Lakhs Only) paid by the VENDEE to the VENDOR -- receipt of which sum is hereby specifically acknowledged by the VENDOR -- the VENDOR hereby absolutely SELLS, TRANSFERS, and CONVEYS, by way of absolute, irrevocable, and indefeasible sale, unto the VENDEE, his heirs, executors, administrators, legal representatives, and assigns:',
   bold_phrases=["absolutely SELLS, TRANSFERS, and CONVEYS", "absolute, irrevocable, and indefeasible sale"],
   space_b=2, space_a=4)

sp(doc, 'ALL THAT piece and parcel of land being Plot No. [.] in the layout on Survey No.240/3A, more particularly described in the SCHEDULE OF PROPERTY (LAND) and SCHEDULE OF PLOT hereunder, together with all rights, easements, liberties, privileges, and appurtenances whatsoever to the said property belonging or in anywise appertaining,',
   bold_phrases=["Plot No. [.]", "SCHEDULE OF PROPERTY (LAND)", "SCHEDULE OF PLOT"],
   space_b=2, space_a=4)

sp(doc, 'TO HAVE AND TO HOLD the same unto the VENDEE, his heirs, executors, administrators, legal representatives, and assigns absolutely and forever, as the sole and absolute owner thereof.',
   bold_phrases=["TO HAVE AND TO HOLD", "sole and absolute owner"], space_b=2, space_a=6)

# Clause 2
heading(doc, "CLAUSE 2 -- VENDOR'S REPRESENTATIONS AND WARRANTIES", size=12, space_b=10, space_a=4)

sp(doc, 'The VENDOR hereby represents and warrants to the VENDEE, as on the date of execution of this Deed, as follows. These representations and warranties are material and form the basis of this transaction, and shall survive the execution, registration, and delivery of this Deed.',
   bold_phrases=["REPRESENTATIONS AND WARRANTIES", "material", "survive"],
   space_b=2, space_a=4)

vendor_reps = [
    ("(i)  Title and Ownership:", "The VENDOR has a valid, absolute, marketable, and indefeasible title to the schedule property, and has full right, title, authority, and power to sell, transfer, and convey the same to the VENDEE by way of absolute sale. The VENDOR is the absolute owner in possession of the schedule property, and no other person or entity has any right, title, interest, or claim over the schedule property."),
    ("(ii)  No Encumbrances:", "The schedule property is free from all encumbrances, mortgages, charges, liens, hypothecations, attachments, deposits, caveats, or demands of any sort, whether known or unknown, and the VENDOR has not created, granted, or suffered any third-party rights over the same."),
    ("(iii)  No Pending or Threatened Litigation:", "The schedule property is not the subject matter of any pending, threatened, or anticipated litigation, arbitration, dispute, claim, prosecution, attachment, or proceedings before any court, tribunal, authority, or forum, except as fully and frankly disclosed in writing by the VENDOR to the VENDEE prior to the execution of this Deed. The existence of O.S. 100/2025 pending before the Hon'ble Additional District Judge, Court Hosur (which relates to the broader title chain of the land), has been disclosed to the VENDEE, and the VENDEE has agreed to proceed with this transaction with full knowledge of the same."),
    ("(iv)  No Acquisition or Requisition:", "The schedule property is not notified or subject to any acquisition or requisition proceedings under the Land Acquisition Act, 1894, or any other central or state law. The VENDOR does not hold the schedule property under the provisions of the Urban Land Ceiling Act, 1976 (Tamil Nadu)."),
    ("(v)  Layout Approval and Compliance:", "The layout in which the schedule property is situated has received full and unconditional Technical Approval from the Hosur New Town Development Authority under Approval Nos. SWP Vz: 90/2025 and 38/2025. All conditions of such approval have been complied with by the VENDOR, and no condition remains outstanding or has been waived."),
    ("(vi)  RERA Compliance:", "The VENDOR/Promoter has complied, and/or undertakes to comply in full, with all applicable requirements of the Real Estate (Regulation and Development) Act, 2016 and the Tamil Nadu Real Estate Regulatory Authority (TNRERA), including registration of the project with TNRERA, in accordance with G.O.Ms.No.112 dated 22.06.2017, before and/or concurrent with the execution and registration of this Sale Deed."),
    ("(vii)  Gift Deeds for Roads and Common Areas:", "The VENDOR has executed registered Gift Deeds for all road areas and common areas within the layout in favour of the relevant government authorities -- Gift Deed Doc No.1634/2025 (TANGEDCO) and Gift Deed Doc No.1632/2025 (Sevaganapalli Panchayat) -- both dated 24.02.2025, and the same are on record in due form."),
    ("(viii)  No Government Dues or Statutory Defaults:", "All property taxes, betterment charges, improvement trust dues, water charges, electricity charges, statutory fees, and costs pertaining to the schedule property have been paid in full up to the date of execution of this Deed. No statutory dues, interest, or penalties are outstanding against the schedule property."),
    ("(ix)  Land Use Conversion:", "The schedule property has been validly converted from agricultural use to non-agricultural/residential use in accordance with applicable Tamil Nadu land conversion laws and rules. The layout is approved by the competent authority, and the schedule property is a lawful residential plot within the approved layout."),
    ("(x)  No Breach of Agreement with Prior Purchasers:", "The VENDOR has not entered into any agreement, MoU, or arrangement with any prior purchaser, allottee, or third party in respect of the schedule property or any part thereof that remains unfulfilled or that would affect the VENDEE's title or possession. All prior agreements pertaining to plots within the layout have been fully performed or duly cancelled, and no claims are outstanding from any prior purchaser."),
    ("(xi)  Force Majeure:", "To the best of the VENDOR's knowledge and belief, no event of force majeure, act of God, or government restriction exists or is anticipated that would materially affect the VENDOR's ability to convey clear title to the schedule property."),
    ("(xii)  TDS / Income Tax Compliance:", "The VENDOR and VENDEE confirm that the total sale consideration herein is Rs.24,00,000/- (Rupees Twenty Four Lakhs Only), which is below the Rs.50,00,000/- (Rupees Fifty Lakhs Only) threshold specified under Section 194-IA of the Income Tax Act, 1961. Accordingly, TDS deduction under Section 194-IA is not applicable to this transaction. Should the applicable threshold or TDS requirements change, or should any correction or demand be raised by the Income Tax authorities in relation to this transaction, the VENDEE undertakes to comply with all applicable TDS obligations and provide the VENDOR with such TDS certificates as may be required."),
]

for roman, text in vendor_reps:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = p.paragraph_format
    pf.left_indent = Inches(0.5)
    pf.first_line_indent = Inches(-0.3)
    pf.space_before = Pt(3)
    pf.space_after = Pt(3)
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    r1 = p.add_run(roman + "  "); r1.bold = True; r1.font.size = Pt(11)
    r2 = p.add_run(text); r2.font.size = Pt(11)

# Clause 3
heading(doc, "CLAUSE 3 -- VENDOR'S COVENANTS", size=12, space_b=10, space_a=4)

sp(doc, 'The VENDOR hereby further covenants with the VENDEE as follows:',
   bold_phrases=["VENDOR'S COVENANTS"], space_b=2, space_a=3)

vendor_covenants = [
    ("(i)  Further Assurance:", "The VENDOR and all persons claiming under the VENDOR shall and will, from time to time and at all times hereafter, upon the reasonable request and at the cost and expense of the VENDEE, do, execute, and cause to be done all such acts, deeds, documents, and things as may be reasonably required to perfect, confirm, and complete the title and possession of the VENDEE over the schedule property."),
    ("(ii)  Indemnity:", "The VENDOR agrees and undertakes to indemnify, defend, and hold harmless the VENDEE against all losses, damages, costs, charges, and expenses of any nature whatsoever that the VENDEE may suffer, incur, or become liable for, arising from or on account of: (a) any defect in the title of the VENDOR to the schedule property; (b) any breach of the representations and warranties made by the VENDOR under this Deed; or (c) any claim from any prior purchaser, allottee, or third party in respect of the schedule property."),
    ("(iii)  Physical Possession and Enjoyment:", "The VENDOR hereby confirms that absolute, full, and actual physical possession of the schedule property has been delivered to the VENDEE prior to or on the date of execution of this Deed. The VENDEE is entitled to collect, enjoy, and appropriate all rents, profits, and benefits arising from the schedule property from the date of execution hereof. The VENDOR shall take responsibility for rectifying any defects in the title, upon notice from the VENDEE, within a reasonable period."),
    ("(iv)  Title Defence:", "The VENDOR shall defend the VENDEE's title and possession against any person or entity claiming through or under the VENDOR, at the VENDOR's own cost and expense, and shall execute all documents as may be reasonably required to maintain the VENDEE's title to the schedule property."),
]

for roman, text in vendor_covenants:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = p.paragraph_format
    pf.left_indent = Inches(0.5)
    pf.first_line_indent = Inches(-0.3)
    pf.space_before = Pt(3)
    pf.space_after = Pt(3)
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    r1 = p.add_run(roman + "  "); r1.bold = True; r1.font.size = Pt(11)
    r2 = p.add_run(text); r2.font.size = Pt(11)

# Clause 4
heading(doc, "CLAUSE 4 -- VENDEE'S ACKNOWLEDGMENTS AND COVENANTS", size=12, space_b=10, space_a=4)

sp(doc, 'The VENDEE hereby acknowledges and covenants with the VENDOR as follows. These acknowledgments are material and form the basis of this transaction:',
   bold_phrases=["VENDEE'S ACKNOWLEDGMENTS AND COVENANTS", "material"],
   space_b=2, space_a=3)

vendee_items = [
    ("(i)  Review of Title Documents:", "The VENDEE acknowledges that he has, prior to the execution of this Deed, carefully reviewed, examined, and scrutinised all title documents pertaining to the schedule property through his own legal advisors, advocates, and counsel of his choice. The VENDEE confirms that he has been fully satisfied as to the validity, marketability, and completeness of the VENDOR's title to the schedule property."),
    ("(ii)  Review of Legal Documentation:", "The VENDEE confirms that he has reviewed all legal documentation pertaining to the layout and the schedule property, including but not limited to the layout approval, RERA compliance documents, Gift Deeds for roads and common areas, Encumbrance Certificate, and all other documents pertaining to the chain of title, either personally or through his legal counsel, and has found the same to be in order."),
    ("(iii)  Independent Legal Advice:", "The VENDEE confirms that he has had the opportunity to obtain independent legal advice from his own advocate and/or counsel before executing this Deed, and has either availed such advice or consciously chosen not to do so, in either case with full knowledge of the implications of this Deed."),
    ("(iv)  Physical Inspection, Measurement, and Satisfaction:", "The VENDEE confirms that he has physically inspected the schedule property, measured its dimensions, and verified the extent and boundaries of the schedule property. The VENDEE is fully satisfied with the physical condition, dimensions, boundaries, extent, and possession of the schedule property, and has agreed to purchase the same of his own free will and accord. The VENDEE acknowledges having received and reviewed the layout plan, approved by HNTDA, showing the location and dimensions of the schedule property within the layout."),
    ("(v)  Registration in Vendee's Own Name:", "The VENDEE is registering/has registered this Deed in his own name in the Sub-Registrar's Office, Bagalur, Tamil Nadu, as the absolute owner of the schedule property. The VENDEE takes full responsibility for ensuring timely presentation and registration of this Deed."),
    ("(vi)  No Sole Reliance on Vendor's Representations:", "The VENDEE acknowledges that he has not relied solely on the representations and warranties of the VENDOR in entering into this transaction, but has conducted his own due diligence and independent verification of all matters material to this transaction through his own legal counsel."),
    ("(vii)  Stamp Duty and Registration Costs:", "The VENDEE shall bear and pay all costs, charges, and expenses towards stamp duty, registration fees, and other incidental expenses for and in connection with the execution, presentation, and registration of this Deed before the Sub-Registrar, Bagalur, Tamil Nadu."),
    ("(viii)  No Additional Claims:", "The VENDEE confirms that upon execution and registration of this Deed and delivery of physical possession, he shall have no further claims, demands, or dues of any nature whatsoever against the VENDOR in respect of the schedule property. The VENDEE acknowledges that the entire sale consideration has been paid and received, and no amount remains outstanding from either party."),
]

for roman, text in vendee_items:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = p.paragraph_format
    pf.left_indent = Inches(0.5)
    pf.first_line_indent = Inches(-0.3)
    pf.space_before = Pt(3)
    pf.space_after = Pt(3)
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    r1 = p.add_run(roman + "  "); r1.bold = True; r1.font.size = Pt(11)
    r2 = p.add_run(text); r2.font.size = Pt(11)

# Clause 5
heading(doc, "CLAUSE 5 -- CONSIDERATION", size=12, space_b=10, space_a=4)

cons_items = [
    ("(i)  Total Sale Consideration:", "The total sale consideration for the schedule property is Rs.24,00,000/- (Rupees Twenty Four Lakhs Only)."),
    ("(ii)  Payment and Receipt:", "The VENDEE has paid, and the VENDOR has received, the entire sale consideration of Rs.24,00,000/-, prior to or simultaneously with the execution of this Deed. The VENDOR acknowledges receipt of the same in full and confirms that no amount is outstanding from the VENDEE."),
    ("(iii)  Declaration of Market Value:", "The VENDOR and VENDEE hereby jointly declare the present market value of the schedule property as Rs.24,00,000/- (Rupees Twenty Four Lakhs Only) for the purposes of registration of this Deed."),
    ("(iv)  Mode of Payment -- Demand Draft:", "The entire sale consideration of Rs.24,00,000/- (Rupees Twenty Four Lakhs Only) has been paid by the VENDEE to the VENDOR by way of a Demand Draft bearing No.[.], drawn on [Bank Name], [Branch], issued on [Date]. The VENDOR confirms receipt of the said Demand Draft and acknowledges that the proceeds have been realised and accepted by the VENDOR. The VENDOR shall deliver the original Demand Draft to the VENDEE at the time of registration of this Deed."),
    ("(v)  Confirmation of Payment at Registration:", "Both VENDOR and VENDEE confirm that the full sale consideration of Rs.24,00,000/- (Rupees Twenty Four Lakhs Only) stands fully paid and received prior to the execution of this Deed, and that no further sum is outstanding from the VENDEE to the VENDOR. The VENDOR confirms readiness to register this Deed upon receipt of the Demand Draft proceeds."),
]

for roman, text in cons_items:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = p.paragraph_format
    pf.left_indent = Inches(0.5)
    pf.first_line_indent = Inches(-0.3)
    pf.space_before = Pt(3)
    pf.space_after = Pt(3)
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    r1 = p.add_run(roman + "  "); r1.bold = True; r1.font.size = Pt(11)
    r2 = p.add_run(text); r2.font.size = Pt(11)

# Clause 6
heading(doc, "CLAUSE 6 -- SCHEDULE OF PROPERTY (LAND DESCRIPTION)", size=12, space_b=10, space_a=4)

sp(doc, 'All that piece and parcel of land situated, lying, and being at:',
   bold_phrases=["SCHEDULE OF PROPERTY (LAND DESCRIPTION)"], space_b=2, space_a=3)

kv(doc, "Village:", "Sevaganapalli")
kv(doc, "Taluk:", "Hosur")
kv(doc, "District:", "Krishnagiri")
kv(doc, "Sub-Registration District:", "Bagalur")
kv(doc, "State:", "Tamil Nadu")
kv(doc, "Original Survey No.:", "240/3")
kv(doc, "Sub-divided Survey No.:", "240/3A (as mutated under Joint Patta No.1665)")
kv(doc, "Original Extent:", "Hec.0.74.0 (Ac.1.83 cents)")
kv(doc, "Sub-divided Extent:", "Hec.0.70.76 (Ac.1.75 cents)")

sp(doc, 'The schedule property forms part of the residential plot layout developed on Survey No.240/3A, Sevaganapalli Village, Hosur Taluk, Krishnagiri District, Tamil Nadu, as approved by the Hosur New Town Development Authority under Approval Nos. SWP Vz: 90/2025 and 38/2025.',
   bold_phrases=["Survey No.240/3A", "Hosur New Town Development Authority"],
   space_b=4, space_a=6)

# Clause 7
heading(doc, "CLAUSE 7 -- SCHEDULE OF PLOT", size=12, space_b=10, space_a=4)

kv(doc, "Plot No.:", "[.]  (to be filled at registration)")
kv(doc, "Layout:", "Residential Plot Layout on Survey No.240/3A, Sevaganapalli Village, Hosur Taluk, Krishnagiri District")
kv(doc, "Total Plot Area:", "111.41 square metres (approximately 1,200 square feet)")

sp(doc, 'Dimensions:', bold_phrases=["Dimensions:"], space_b=3, space_a=2)
sub_para(doc, "East to West -- Northern side: 9.14 metres")
sub_para(doc, "East to West -- Southern side: 9.14 metres")
sub_para(doc, "North to South -- Eastern side: 12.19 metres")
sub_para(doc, "North to South -- Western side: 12.19 metres")

sp(doc, 'Boundaries:', bold_phrases=["Boundaries:"], space_b=4, space_a=2)
kv(doc, "East:", "Land in Survey No.240/2A2")
kv(doc, "West:", "Road")
kv(doc, "North:", "Plot No.[.]  (to be filled at registration)")
kv(doc, "South:", "Plot No.[.]  (to be filled at registration)")

sp(doc, 'Location: Within the limits of Sevaganapalli Panchayat and Union Council of Hosur, Krishnagiri District, Tamil Nadu.',
   bold_phrases=["Sevaganapalli Panchayat", "Union Council of Hosur"],
   space_b=4, space_a=6)

# Clause 8
heading(doc, "CLAUSE 8 -- GENERAL PROVISIONS", size=12, space_b=10, space_a=4)

gen_provisions = [
    ("(i)", "The headings in this Deed are for convenience of reference only and shall not affect the interpretation or construction of this Deed."),
    ("(ii)", "The Schedules to this Deed shall form an integral part of this Deed."),
    ("(iii)", "Any waiver by either Party of any breach or default by the other Party shall not be deemed a waiver of any subsequent breach or default, nor shall any waiver constitute a continuing waiver."),
    ("(iv)", "This Deed shall be governed by and construed in accordance with the laws of India. The courts at Krishnagiri, Tamil Nadu, shall have exclusive jurisdiction over any disputes arising under or in connection with this Deed."),
    ("(v)", "Words importing the singular shall include the plural and vice versa. Words importing a gender shall include all genders."),
    ("(vi)", "If any provision of this Deed is held to be invalid, illegal, or unenforceable by a court of competent jurisdiction, such invalidity, illegality, or unenforceability shall not affect the validity, legality, or enforceability of the remaining provisions of this Deed, which shall continue in full force and effect."),
]

for roman, text in gen_provisions:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = p.paragraph_format
    pf.left_indent = Inches(0.5)
    pf.first_line_indent = Inches(-0.3)
    pf.space_before = Pt(3)
    pf.space_after = Pt(3)
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    r1 = p.add_run(roman + "  "); r1.bold = True; r1.font.size = Pt(11)
    r2 = p.add_run(text); r2.font.size = Pt(11)

# Clause 9
heading(doc, "CLAUSE 9 -- TESTIMONIUM AND EXECUTION", size=12, space_b=10, space_a=4)

sp(doc, 'IN WITNESS WHEREOF, the VENDOR and the VENDEE have signed this Deed of Absolute Sale on the day, month, and year first above written, at Bagalur, Krishnagiri District, Tamil Nadu.',
   bold_phrases=["IN WITNESS WHEREOF"], space_b=2, space_a=8)

heading(doc, "EXECUTED BY THE VENDOR:", size=11, space_b=4, space_a=3)

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(4)
p.paragraph_format.space_after = Pt(2)
p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
r1 = p.add_run("For and on behalf of M/s. DRA THINDLU LAND PARTNERS"); r1.bold = True; r1.font.size = Pt(11)

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(2)
p.paragraph_format.space_after = Pt(1)
p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
p.add_run("Partner 1 -- M/s. DRA REALTY PRIVATE LIMITED (Director: Mr. Kishan Murjani Nair)").font.size = Pt(11)

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(1)
p.paragraph_format.space_after = Pt(1)
p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
p.add_run("Signature: ___________________________________    Name: KISHAN MURJANI NAIR    Aadhaar: 970838057634").font.size = Pt(11)

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(4)
p.paragraph_format.space_after = Pt(1)
p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
p.add_run("Partner 2 -- Mr. NISHANT RANKA (Partner, DRA Thindlu Land Partners)").font.size = Pt(11)

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(1)
p.paragraph_format.space_after = Pt(1)
p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
p.add_run("Signature: ___________________________________    Name: NISHANT RANKA    Aadhaar: 4159 0535 2796").font.size = Pt(11)

heading(doc, "EXECUTED BY THE VENDEE:", size=11, space_b=6, space_a=3)

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(2)
p.paragraph_format.space_after = Pt(1)
p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
r1 = p.add_run("Mr. MANJUNATH SINGH MANOHAR SINGH"); r1.bold = True; r1.font.size = Pt(11)

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(1)
p.paragraph_format.space_after = Pt(1)
p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
p.add_run("Signature: ___________________________________    Name: MANJUNATH SINGH MANOHAR SINGH    PAN: AOQPS1456J    Aadhaar: 4480 4783 5557").font.size = Pt(11)

heading(doc, "WITNESSES:", size=11, space_b=6, space_a=3)

for w_num in ["1", "2"]:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.add_run(f"Name: ___________________________________    Address: ___________________________________    Aadhaar: ___________________________________    Signature: ___________________________________").font.size = Pt(11)

sp(doc, 'PLACE: Bagalur, Krishnagiri District, Tamil Nadu     |     DATE: _____ May 2026',
   bold_phrases=["PLACE:", "Bagalur", "DATE:", "_____ May 2026"],
   space_b=8, space_a=4, align=WD_ALIGN_PARAGRAPH.LEFT)

sp(doc, 'DRAFTED BY: [Advocate / Document Writer Name]  |  [Enrollment / Registration No.]  |  [Full Address]  |  [Contact Number]',
   bold_phrases=["DRAFTED BY:"], space_b=2, space_a=2)

divider(doc)

# Annexures
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(10)
p.paragraph_format.space_after = Pt(4)
r = p.add_run("ANNEXURES"); r.bold = True; r.font.size = Pt(13)

sp(doc, 'The following documents are attached to and form part of this Sale Deed. All other documents referenced in the Background and Recitals form part of the title chain and are available in the original registered form at the relevant Sub-Registrar\'s Office.',
   space_b=2, space_a=4)

annexure_items = [
    ("Annexure A", "Layout Approval Plan -- Hosur New Town Development Authority (HNTDA) Technical Approval No. SWP Vz: 90/2025 and 38/2025, showing the layout of the land on Survey No.240/3A, Sevaganapalli Village, with the schedule property (Plot No.[.]) marked and dimensioned."),
    ("Annexure B", "Revenue Records / Latest Encumbrance Certificate -- issued by the Sub-Registrar, Bagalur, covering the search period from 01.01.1975 to the date of execution of this Deed, showing the schedule property to be free from all encumbrances, charges, and liabilities."),
    ("Annexure C", "Copy of Registered Sale Deed Doc No.20527/2024 -- the VENDOR\'s source title deed, dated 24.10.2024, registered at the Sub-Registrar\'s Office, Hosur, evidencing the VENDOR\'s acquisition of the schedule property."),
]

for annex, desc in annexure_items:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    r1 = p.add_run(annex + ":  "); r1.bold = True; r1.font.size = Pt(11)
    r2 = p.add_run(desc); r2.font.size = Pt(11)

# Save
doc.save('/tmp/sale_deed_v3.docx')
print("Saved. Paragraphs:", len(doc.paragraphs))