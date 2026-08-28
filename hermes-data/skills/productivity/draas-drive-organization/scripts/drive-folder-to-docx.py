#!/usr/bin/env python3
"""Generate a Word (docx) list of all documents in a Google Drive folder tree.
One table per folder. File names are clickable Drive hyperlinks. Ends with a
Summary table. Works for land/legal doc folders (e.g. Oasis - print).

Usage (terminal ONLY — the execute_code sandbox has no vault socket):
  cd /opt/hermes && HERMES_SESSION_USER_ID=psingh /opt/hermes/.venv/bin/python3 \
    /data/hermes/skills/productivity/draas-drive-organization/scripts/drive-folder-to-docx.py \
    FOLDER_ID /tmp/out.docx [SERVICE_NAME] [ROOT_LABEL]

Set HERMES_SESSION_USER_ID to the canonical uid (e.g. psingh for Prakash's account).
Prints the authenticated account + per-folder counts + saved path.
"""
import sys, json, os, datetime
sys.path.insert(0, '/opt/hermes')
from tools.gws_auth import _load_credentials_direct
from googleapiclient.discovery import build
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.opc.constants import RELATIONSHIP_TYPE as RT

FOLDER_ID = sys.argv[1] if len(sys.argv) > 1 else None
OUT = sys.argv[2] if len(sys.argv) > 2 else '/tmp/Drive_Document_List.docx'
SERVICE = sys.argv[3] if len(sys.argv) > 3 else 'google-draas'
ROOT_LABEL = sys.argv[4] if len(sys.argv) > 4 else 'Drive folder'
if not FOLDER_ID:
    sys.exit("Pass FOLDER_ID as argv[1]")

creds = _load_credentials_direct(SERVICE)
svc = build('drive', 'v3', credentials=creds)
me = svc.about().get(fields='user').execute()
print("ACCOUNT:", me['user']['emailAddress'])

def list_children(fid):
    out, token = [], None
    while True:
        r = svc.files().list(q=f"'{fid}' in parents and trashed=false",
                             fields='nextPageToken, files(id,name,mimeType,modifiedTime,size,webViewLink)',
                             pageSize=1000, pageToken=token, supportsAllDrives=True).execute()
        out.extend(r.get('files', []))
        token = r.get('nextPageToken')
        if not token:
            break
    return out

structure = []
def walk(fid, path):
    items = list_children(fid)
    files = [i for i in items if i['mimeType'] != 'application/vnd.google-apps.folder']
    subf = [i for i in items if i['mimeType'] == 'application/vnd.google-apps.folder']
    files.sort(key=lambda x: x['name'].lower())
    subf.sort(key=lambda x: x['name'].lower())
    structure.append({'path': path, 'id': fid, 'name': path.split('/')[-1], 'files': files})
    for s in subf:
        walk(s['id'], path + '/' + s['name'])

root_meta = svc.files().get(fileId=FOLDER_ID, fields='name').execute()
root_name = root_meta.get('name', ROOT_LABEL)
walk(FOLDER_ID, root_name)
print("FOLDERS:", len(structure), "| TOTAL FILES:", sum(len(s['files']) for s in structure))

# ---------- docx rendering ----------
MIME_LABEL = {
    'application/pdf': 'PDF', 'image/jpeg': 'JPEG', 'image/png': 'PNG',
    'application/vnd.openxmlformats-officedocument.wordprocessingml.document': 'DOCX',
    'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet': 'XLSX',
    'application/vnd.google-apps.document': 'Google Doc',
    'application/vnd.google-apps.spreadsheet': 'Google Sheet',
    'application/vnd.google-apps.folder': 'Folder',
}
def mime_label(m):
    return MIME_LABEL.get(m, m.split('.')[-1].upper() if '.' in m else m)

def fmt_date(s):
    try:
        return datetime.datetime.fromisoformat(str(s).replace('Z', '+00:00')).strftime('%d-%b-%Y')
    except Exception:
        return s or ''

def set_cell_bg(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), hexcolor)
    tcPr.append(shd)

def make_table(doc, files):
    t = doc.add_table(rows=len(files) + 1, cols=4)
    t.style = 'Table Grid'
    widths = [Inches(0.55), Inches(4.35), Inches(1.5), Inches(1.6)]
    for j, h in enumerate(['Sl No', 'Document Name', 'Type', 'Modified Date']):
        cell = t.cell(0, j); cell.text = h
        set_cell_bg(cell, 'D9E2F3')
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.font.bold = True; r.font.size = Pt(10)
    for i, f in enumerate(files, 1):
        vals = [str(i), f['name'], mime_label(f['mimeType']), fmt_date(f.get('modifiedTime'))]
        for j, v in enumerate(vals):
            cell = t.rows[i].cells[j]; cell.text = v
            for p in cell.paragraphs:
                p.paragraph_format.space_before = Pt(1)
                p.paragraph_format.space_after = Pt(1)
                for r in p.runs:
                    r.font.size = Pt(9)
        # hyperlink on the document-name cell
        url = f.get('webViewLink') or f"https://drive.google.com/file/d/{f['id']}/view"
        rel_id = doc.part.relate_to(url, RT.HYPERLINK, is_external=True)
        p = t.rows[i].cells[1].paragraphs[0]
        hyperlink = OxmlElement('w:hyperlink'); hyperlink.set(qn('r:id'), rel_id)
        new_run = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')
        color = OxmlElement('w:color'); color.set(qn('w:val'), '0563C1'); rPr.append(color)
        u = OxmlElement('w:u'); u.set(qn('w:val'), 'single'); rPr.append(u)
        sz = OxmlElement('w:sz'); sz.set(qn('w:val'), '18'); rPr.append(sz)
        new_run.append(rPr)
        t_el = OxmlElement('w:t'); t_el.text = f['name']; t_el.set(qn('xml:space'), 'preserve')
        new_run.append(t_el); hyperlink.append(new_run)
        for child in list(p._p):  # clear default run, replace with hyperlink
            p._p.remove(child)
        p._p.append(hyperlink)
    for row in t.rows:  # honour widths on every row
        for j, w in enumerate(widths):
            row.cells[j].width = w

doc = Document()
style = doc.styles['Normal']; style.font.name = 'Calibri'; style.font.size = Pt(10)
title = doc.add_heading(f'{root_name} — Document List', level=0)
for r in title.runs:
    r.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)
sub = doc.add_paragraph()
r = sub.add_run(f'Drive folder: {root_name}\nGenerated by Hermes | {datetime.date.today().isoformat()} | Account: {me["user"]["emailAddress"]}')
r.font.size = Pt(9); r.font.color.rgb = RGBColor(0x59, 0x59, 0x59)

grand_total = 0
for idx, st in enumerate(structure, 1):
    doc.add_heading(f"Folder {idx}: {st['path']}", level=1)
    p = doc.add_paragraph()
    r = p.add_run(f"Documents: {len(st['files'])}")
    r.font.size = Pt(9); r.font.italic = True; r.font.color.rgb = RGBColor(0x59, 0x59, 0x59)
    if st['files']:
        make_table(doc, st['files'])
    else:
        p2 = doc.add_paragraph(); r2 = p2.add_run('(empty folder)'); r2.font.italic = True
    grand_total += len(st['files'])

doc.add_heading('Summary', level=1)
sum_tbl = doc.add_table(rows=len(structure) + 2, cols=2)
sum_tbl.style = 'Table Grid'
for j, h in enumerate(['Folder', 'Documents']):
    c = sum_tbl.rows[0].cells[j]; c.text = h
    set_cell_bg(c, 'D9E2F3')
    for p in c.paragraphs:
        for r in p.runs: r.font.bold = True
for i, st in enumerate(structure, 1):
    sum_tbl.rows[i].cells[0].text = st['path']
    sum_tbl.rows[i].cells[1].text = str(len(st['files']))
sum_tbl.rows[len(structure) + 1].cells[0].text = 'TOTAL'
sum_tbl.rows[len(structure) + 1].cells[1].text = str(grand_total)
for row in sum_tbl.rows:
    row.cells[0].width = Inches(5.5); row.cells[1].width = Inches(1.5)

doc.save(OUT)
print("SAVED:", OUT, "| folders:", len(structure), "| total files:", grand_total)