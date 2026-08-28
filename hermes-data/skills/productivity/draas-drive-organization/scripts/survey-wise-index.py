#!/usr/bin/env python3
"""
Generate a survey-no.-wise Word document index for any Drive folder.

Usage:
    cd /opt/hermes && HERMES_SESSION_USER_ID=psingh \\
      /opt/hermes/.venv/bin/python skills/productivity/draas-drive-organization/scripts/survey-wise-index.py \\
      <FOLDER_ID> <output.docx> [service_name]

Args:
    FOLDER_ID   — Google Drive folder ID
    output.docx — path for the output Word file
    service_name — GWS vault key (default: google-draas)

The script:
  1. Walks ALL subfolders recursively
  2. Extracts survey numbers from filenames using Indian real-estate patterns
  3. Groups documents survey-no.-wise, sorted oldest→newest
  4. Remaining documents grouped by folder name
  5. Writes a Word doc with tables
"""
import sys, json, re, os
from datetime import datetime

sys.path.insert(0, '/opt/hermes')
from tools.gws_auth import _load_credentials_direct
from googleapiclient.discovery import build
from googleapiclient.http import MediaIoBaseDownload
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import io

FOLDER_ID = sys.argv[1] if len(sys.argv) > 1 else None
OUT_PATH = sys.argv[2] if len(sys.argv) > 2 else '/tmp/survey_wise_index.docx'
SERVICE = sys.argv[3] if len(sys.argv) > 3 else 'google-draas'

if not FOLDER_ID:
    print("Usage: survey-wise-index.py <FOLDER_ID> <output.docx> [service_name]")
    sys.exit(1)

# ── Auth ──
creds = _load_credentials_direct(SERVICE)
drive = build('drive', 'v3', credentials=creds)
user = drive.about().get(fields='user').execute()
print(f"Authenticated as: {user['user']['emailAddress']}")

# ── Walk folder ──
def walk_folder(fid, path):
    items, token = [], None
    while True:
        r = drive.files().list(
            q=f"'{fid}' in parents and trashed=false",
            fields='nextPageToken, files(id,name,mimeType,size,webViewLink)',
            pageSize=1000, pageToken=token, supportsAllDrives=True
        ).execute()
        items.extend(r.get('files', []))
        token = r.get('nextPageToken')
        if not token:
            break
    files = [i for i in items if i['mimeType'] != 'application/vnd.google-apps.folder']
    subs = [i for i in items if i['mimeType'] == 'application/vnd.google-apps.folder']
    result = {'path': path, 'name': path.split('/')[-1], 'files': files}
    for s in sorted(subs, key=lambda x: x['name'].lower()):
        result = result  # no-op, collect into flat list
    all_folders = [result]
    for s in sorted(subs, key=lambda x: x['name'].lower()):
        all_folders.extend(walk_folder(s['id'], f"{path}/{s['name']}"))
    return all_folders

print("Walking folder tree...")
structure = walk_folder(FOLDER_ID, 'root')
total_files = sum(len(f['files']) for f in structure)
print(f"Found {len(structure)} folders, {total_files} files")

# ── Extraction ──
def extract_date(name):
    n = name.strip()
    n = re.sub(r'^Copy of\s+', '', n, flags=re.I)
    m = re.match(r'^(\d{8})(?!\d)', n)
    if m:
        s = m.group(1)
        y, mo, d = int(s[:4]), int(s[4:6]), int(s[6:8])
        if 1900 <= y <= 2100 and 1 <= mo <= 12 and 1 <= d <= 31:
            return datetime(y, mo, d)
        dd, mm, yyyy = int(s[:2]), int(s[2:4]), int(s[4:])
        if 1900 <= yyyy <= 2100 and 1 <= mm <= 12 and 1 <= dd <= 31:
            return datetime(yyyy, mm, dd)
    m = re.search(r'(?:dtd\.?|dated?)\s+(\d{1,2})[-./](\d{1,2})[-./](\d{4})', n, re.I)
    if m:
        d, mo, y = int(m.group(1)), int(m.group(2)), int(m.group(3))
        if 1 <= d <= 31 and 1 <= mo <= 12 and 1900 <= y <= 2100:
            return datetime(y, mo, d)
    return None

def norm_sn(s):
    s = s.strip('.()').strip()
    s = re.sub(r'[-_]', '/', s)
    parts = s.split('/')
    return '/'.join([re.sub(r'^0+(?=\d)', '', p.strip()) for p in parts])

def extract_surveys(name):
    results = set()
    for m in re.finditer(r'(?:sy[.\s]*no[.\s]*\'?s?[:\s\-]*)(\(?\s*[A-Za-z0-9/.\-_]+\s*\)?)', name, re.I):
        raw = m.group(1).strip().rstrip(',;.)')
        for part in re.split(r'[,&]\s+|\s+and\s+', raw):
            part = part.strip().rstrip(',;.)(')
            if part and re.search(r'\d', part) and not re.match(r'^\d+$', part):
                results.add(norm_sn(part))
    for m in re.finditer(r'\bFMB\s+(\d+(?:[.\-/(]\d+[A-Za-z0-9]*)*)', name, re.I):
        results.add(norm_sn(m.group(1).rstrip(')')))
    for m in re.finditer(r'\bPatta\s+(?:no\.?\s*)?(\d+\s*\(?[A-Za-z0-9]*\)?)', name, re.I):
        results.add(norm_sn(m.group(1).rstrip(')')))
    for m in re.finditer(r'\bUDR[:\s]+(?:SY\s*NO\s+)?(\d+(?:[.\-/\s]\d+[A-Za-z0-9]*)*)', name, re.I):
        for part in re.split(r'\s*-\s*', m.group(1).strip().rstrip(')')):
            if part.strip() and re.search(r'\d', part):
                results.add(norm_sn(part.strip()))
    for m in re.finditer(r'\bEC\s.*?\b(?:Sy|SY)[.\s]*[Nn][Oo]?[.\s]*(\d+(?:[._/\-]\d+[A-Za-z0-9]*)*)', name, re.I):
        results.add(norm_sn(m.group(1).rstrip(')')))
    for m in re.finditer(r'\b[Aa]dangal\b.*?(\d+(?:[._/\-]\d+[A-Za-z0-9]*)+)', name, re.I):
        results.add(norm_sn(m.group(1).rstrip(')')))
    for m in re.finditer(r'(?:sy[.\s]*no[.\s]*|syno[.\s]*)[.\-]?\s*(\d+(?:[.\-/\s]\d+[A-Za-z0-9]*)*)', name, re.I):
        for part in re.split(r'[,/]\s*', m.group(1).strip().rstrip(')')):
            if part.strip() and re.search(r'\d', part):
                results.add(norm_sn(part.strip()))
    for m in re.finditer(r'(?:Sy|SY)\s+(\d+(?:[/\-]\d+[A-Za-z0-9]*)*(?:,\s*\d+(?:[/\-]\d+[A-Za-z0-9]*)*)*)', name):
        for part in re.split(r',\s*', m.group(1).strip()):
            results.add(norm_sn(part))
    # Filter: patta-only numbers, area values
    skip = {'25', '62', '357', '405', '543', '581', '643', '1006', '1007', '1008', '1009', '1117', '1263', '1393', '1441', '1776', '1922', '2000', '7.22', '25 SY'}
    clean = set()
    for s in results:
        if s in skip or re.match(r'^\d{2,3}$', s):
            continue
        clean.add(s)
    return sorted(clean) if clean else []

def set_cell_bg(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), hexcolor)
    tcPr.append(shd)

# ── Build data ──
survey_docs = {}
folder_docs = {}
for folder in structure:
    fp = folder['path']
    for f in folder['files']:
        entry = {'name': f['name'], 'date': extract_date(f['name']),
                 'date_str': extract_date(f['name']).strftime('%d-%m-%Y') if extract_date(f['name']) else '',
                 'folder': fp, 'id': f.get('id', '')}
        sv = extract_surveys(f['name'])
        if sv:
            for s in sv:
                survey_docs.setdefault(s, []).append(entry)
        else:
            folder_docs.setdefault(fp, []).append(entry)

def sort_key(sn):
    parts = re.split(r'/|\.', sn)
    nums = [int(re.match(r'(\d+)', p).group(1)) if re.match(r'(\d+)', p) else 9999 for p in parts]
    while len(nums) < 3: nums.append(0)
    return tuple(nums)

sorted_surveys = sorted(survey_docs.keys(), key=sort_key)

# ── Build Word doc ──
doc = Document()
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(10)

title = doc.add_heading('Survey-No.-Wise Document Index', level=0)
for r in title.runs:
    r.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)

sub = doc.add_paragraph()
r = sub.add_run(f'Folder ID: {FOLDER_ID}  |  Account: {user["user"]["emailAddress"]}  |  Files: {total_files}')
r.font.size = Pt(9); r.font.color.rgb = RGBColor(0x59, 0x59, 0x59)

doc.add_heading('PART 1 — Survey-No.-Wise Documents', level=1)

total_grouped = 0
for sn in sorted_surveys:
    docs = survey_docs[sn]
    docs.sort(key=lambda x: (x['date'] or datetime.max, x['name'].lower()))
    seen = set()
    unique = []
    for d in docs:
        if d['name'] not in seen:
            seen.add(d['name'])
            unique.append(d)
    doc.add_heading(f'Survey: {sn}  ({len(unique)} docs)', level=2)
    t = doc.add_table(rows=len(unique)+1, cols=4)
    t.style = 'Table Grid'
    for j, h in enumerate(['#', 'Date', 'Document Name', 'Folder']):
        c = t.rows[0].cells[j]; c.text = h
        set_cell_bg(c, 'D9E2F3')
        for p in c.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs: r.font.bold = True; r.font.size = Pt(8)
    for i, d in enumerate(unique, 1):
        for j, v in enumerate([str(i), d['date_str'] or '—', d['name'], d['folder']]):
            c = t.rows[i].cells[j]; c.text = v
            for p in c.paragraphs:
                p.paragraph_format.space_before = Pt(1); p.paragraph_format.space_after = Pt(1)
                for r in p.runs: r.font.size = Pt(8)
    total_grouped += len(unique)

doc.add_heading(f'PART 2 — By Folder ({sum(len(v) for v in folder_docs.values())} docs)', level=1)
for fp in sorted(folder_docs.keys()):
    docs = folder_docs[fp]
    docs.sort(key=lambda x: (x['date'] or datetime.max, x['name'].lower()))
    label = fp
    doc.add_heading(f'{label}  ({len(docs)} docs)', level=2)
    t = doc.add_table(rows=len(docs)+1, cols=3)
    t.style = 'Table Grid'
    for j, h in enumerate(['#', 'Date', 'Document Name']):
        c = t.rows[0].cells[j]; c.text = h
        set_cell_bg(c, 'D9E2F3')
        for p in c.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs: r.font.bold = True; r.font.size = Pt(8)
    for i, d in enumerate(docs, 1):
        for j, v in enumerate([str(i), d['date_str'] or '—', d['name']]):
            c = t.rows[i].cells[j]; c.text = v
            for p in c.paragraphs:
                p.paragraph_format.space_before = Pt(1); p.paragraph_format.space_after = Pt(1)
                for r in p.runs: r.font.size = Pt(8)

doc.add_heading('Summary', level=1)
t = doc.add_table(rows=5, cols=2)
t.style = 'Table Grid'
for i, (k, v) in enumerate([
    ('Total files', str(total_files)),
    ('Survey-matched docs', str(total_grouped)),
    ('By-folder docs', str(sum(len(v) for v in folder_docs.values()))),
    ('Survey groups', str(len(sorted_surveys))),
    ('Folders', str(len(structure)))
]):
    t.rows[i].cells[0].text = k
    t.rows[i].cells[1].text = v
    for j in range(2):
        for p in t.rows[i].cells[j].paragraphs:
            for r in p.runs:
                r.font.size = Pt(10)
                if i == 0: r.font.bold = True

doc.save(OUT_PATH)
print(f"\nSAVED: {OUT_PATH}")
print(f"Survey groups: {len(sorted_surveys)}, Grouped: {total_grouped}, Folder-only: {sum(len(v) for v in folder_docs.values())}")