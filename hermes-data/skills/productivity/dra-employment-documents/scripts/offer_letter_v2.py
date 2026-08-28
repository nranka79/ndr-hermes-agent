#!/usr/bin/env python3
"""
DRA Offer Letter v2 — python-docx generator.
Usage: python3 /data/hermes/skills/productivity/dra-employment-documents/scripts/offer_letter_v2.py

Customise: edit the constants at the top (candidate, role, compensation, etc.)
then run. Output goes to /data/hermes/cron/output/.

The script generates a visually formatted offer letter with:
  - Tables for structured data (Sections 1, 2, 3, 4, 7)
  - Alternating row shading (blue-grey / white)
  - Yellow highlight for Saturday exemption rows
  - Bold section headers in blue
  - Performance pay mechanics in bullet points below comp table
  - Clean opening paragraph (no resume biography)
  - HR Policy sent separately note
"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── CUSTOMISE THESE FOR EACH CANDIDATE ──────────────────────────
CANDIDATE_NAME  = "Sai Neha Vaddadi"
CANDIDATE_EMAIL = "esotericarts.ani@gmail.com"
CANDIDATE_PHONE = "+91 7899398273"
CANDIDATE_ADDR  = "Bangalore, Karnataka – [full address to be confirmed at joining]"

ROLE_TITLE     = "Content Creator"
REPORTING_TO   = "Gowri Singh — Content & Marketing Head"
COORDINATES    = "Performance Marketing team (for paid distribution / boosting of content)"
ENTITY         = "DRA Realty Private Limited  (CIN: U70100KA2011PTC058105)"
LOCATION       = "204–206, Prism Greystone, Cunningham Road, Bengaluru – 560052"
START_DATE     = "04 June 2026"

BASE_PAY       = 33_000
ATTENDANCE      = 4_000
PERF_PAY       = 3_000
TOTAL_MONTHLY   = 40_000

PROBATION_MONTHS = 6
SATURDAY_EXEMPT  = True   # set False for default Saturday clause
MAX_SATURDAYS    = 20

LETTER_DATE    = "03 June 2026"
# ─────────────────────────────────────────────────────────────────

OUTPUT_PATH = "/data/hermes/cron/output/20260603_DRA_SaiNehaVaddadi_OfferLetter_ContentCreator_v2.docx"

doc = Document()

for section in doc.sections:
    section.top_margin    = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin   = Cm(2.2)
    section.right_margin  = Cm(2.2)

# ── Helpers ──────────────────────────────────────────────────────
def sf(run, name='Calibri', size=11, bold=False, color=None, italic=False):
    run.font.name  = name
    run.font.size  = Pt(size)
    run.bold       = bold
    run.italic     = italic
    if color:
        run.font.color.rgb = RGBColor(*color)

def cbg(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)

def h1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(3)
    r = p.add_run(text.upper())
    sf(r, bold=True, size=10.5, color=(0, 51, 102))
    return p

def h2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(2)
    r = p.add_run(text)
    sf(r, bold=True, size=10.5, color=(31, 73, 125))
    return p

def body(doc, text, sa=3):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(sa)
    r = p.add_run(text)
    sf(r, size=10.5)
    return p

def blt(doc, bold_prefix, rest, sa=2):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(sa)
    p.paragraph_format.left_indent  = Cm(0.4)
    if bold_prefix:
        rb = p.add_run(bold_prefix)
        sf(rb, bold=True, size=10.5)
    rn = p.add_run(rest)
    sf(rn, size=10.5)
    return p

def th(tbl, headers, bg='1F497D'):
    row = tbl.rows[0]
    for i, h in enumerate(headers):
        c = row.cells[i]
        c.paragraphs[0].clear()
        p = c.paragraphs[0]
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after  = Pt(3)
        r = p.add_run(h)
        sf(r, bold=True, size=10, color=(255, 255, 255))
        cbg(c, bg)

def tr(tbl, vals, shade=None, bf=False):
    row = tbl.add_row()
    for i, v in enumerate(vals):
        c = row.cells[i]
        p = c.paragraphs[0]
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after  = Pt(2)
        r = p.add_run(v)
        sf(r, size=10.5, bold=(bf and i == 0))
        if shade:
            cbg(c, shade)
    return row

def hline(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot  = OxmlElement('w:bottom')
    bot.set(qn('w:val'),   'single')
    bot.set(qn('w:sz'),    '6')
    bot.set(qn('w:space'), '1')
    bot.set(qn('w:color'), '1F497D')
    pBdr.append(bot)
    pPr.append(pBdr)
    return p

def spc(doc, sz=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(0)
    p.paragraph_format.line_spacing = Pt(sz)
    return p

# ── Letterhead ────────────────────────────────────────────────────
co = doc.add_paragraph()
co.alignment = WD_ALIGN_PARAGRAPH.CENTER
co.paragraph_format.space_before = Pt(0)
co.paragraph_format.space_after  = Pt(2)
r = co.add_run("DRA REALTY PRIVATE LIMITED")
sf(r, bold=True, size=14, color=(31, 73, 125))

reg = doc.add_paragraph()
reg.alignment = WD_ALIGN_PARAGRAPH.CENTER
reg.paragraph_format.space_before = Pt(0)
reg.paragraph_format.space_after  = Pt(1)
r2 = reg.add_run("CIN: U70100KA2011PTC058105  |  GSTIN: 29AAPCS9730H1ZO")
sf(r2, size=8.5, color=(89, 89, 89))

addr = doc.add_paragraph()
addr.alignment = WD_ALIGN_PARAGRAPH.CENTER
addr.paragraph_format.space_before = Pt(0)
addr.paragraph_format.space_after  = Pt(2)
r3 = addr.add_run(
    "Registered Office: 4, Ranka Chambers, 31 Cunningham Road, Bangalore – 560052\n"
    "Operational Office: 204–206, Prism Greystone, Cunningham Road, Bengaluru – 560052")
sf(r3, size=8.5, color=(89, 89, 89))

hline(doc)

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.paragraph_format.space_before = Pt(8)
title.paragraph_format.space_after  = Pt(6)
rt = title.add_run("OFFER OF EMPLOYMENT")
sf(rt, bold=True, size=13, color=(31, 73, 125))

for k, v in [("Date:", LETTER_DATE),
             ("To:", CANDIDATE_NAME),
             ("", CANDIDATE_ADDR),
             ("", f"Email: {CANDIDATE_EMAIL}  |  Phone: {CANDIDATE_PHONE}")]:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(1)
    if k:
        r1 = p.add_run(k + "  ")
        sf(r1, bold=True, size=10.5)
    r2 = p.add_run(v)
    sf(r2, size=10.5)

spc(doc, 6)

# ── Opening ───────────────────────────────────────────────────────
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(0)
p.paragraph_format.space_after  = Pt(4)
r1 = p.add_run("We are pleased to offer you the position of ")
sf(r1, size=10.5)
r2 = p.add_run(ROLE_TITLE)
sf(r2, bold=True, size=10.5)
r3 = p.add_run(
    " at DRA Realty Private Limited. We believe your creative energy, animation craft foundation, "
    "and hunger to learn will significantly strengthen the DRA brand across all our projects and platforms.")
sf(r3, size=10.5)

p2 = doc.add_paragraph()
p2.paragraph_format.space_before = Pt(0)
p2.paragraph_format.space_after  = Pt(6)
r4 = p2.add_run("Please find below the clear terms of this offer:")
sf(r4, size=10.5)

# ════════════════════════════════════════════════
# SECTION 1 — Position, Reporting & Start Date
# ════════════════════════════════════════════════
h1(doc, "1.  Position, Reporting & Start Date")
tbl = doc.add_table(rows=1, cols=2)
tbl.style = 'Table Grid'
tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
th(tbl, ["Detail", "Information"])

data1 = [
    ("Position",           ROLE_TITLE),
    ("Reports to",          REPORTING_TO),
    ("Coordinates with",   COORDINATES),
    ("Entity",             ENTITY),
    ("Work Location",      LOCATION),
    ("Proposed Start Date", START_DATE),
]
for idx, (k, v) in enumerate(data1):
    shade = 'EEF3F8' if idx % 2 == 0 else 'FFFFFF'
    row = tbl.add_row()
    for ci, val in enumerate([k, v]):
        c = row.cells[ci]
        p = c.paragraphs[0]
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after  = Pt(3)
        r = p.add_run(val)
        sf(r, size=10.5, bold=(ci == 0))
        cbg(c, shade)

spc(doc, 6)

# ════════════════════════════════════════════════
# SECTION 2 — Compensation
# ════════════════════════════════════════════════
h1(doc, "2.  Compensation")

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(0)
p.paragraph_format.space_after  = Pt(4)
r1 = p.add_run("Total Monthly Pay Package:  ")
sf(r1, bold=True, size=10.5)
r2 = p.add_run(f"\u20b9{TOTAL_MONTHLY:,} (Rupees Forty Thousand only) per month, all-inclusive")
sf(r2, bold=True, size=10.5, color=(31, 73, 125))

tbl2 = doc.add_table(rows=1, cols=3)
tbl2.style = 'Table Grid'
tbl2.alignment = WD_TABLE_ALIGNMENT.LEFT
th(tbl2, ["Component", "Amount (\u20b9/month)", "Notes"])

comp_data = [
    ("Base Salary",          f"{BASE_PAY:,}",   "Fixed"),
    ("Attendance Allowance", f"{ATTENDANCE:,}",  "Subject to regular attendance & adherence to company policies"),
    ("Performance-Based Pay", f"{PERF_PAY:,}",   "Earned against quarterly KPIs — see note below"),
    ("TOTAL",                f"{TOTAL_MONTHLY:,}", "All-inclusive"),
]
for idx, row_data in enumerate(comp_data):
    is_total = (idx == len(comp_data) - 1)
    shade    = 'D6E4F0' if is_total else ('EEF3F8' if idx % 2 == 0 else 'FFFFFF')
    row = tbl2.add_row()
    for ci, val in enumerate(row_data):
        c = row.cells[ci]
        p = c.paragraphs[0]
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after  = Pt(3)
        r = p.add_run(val)
        sf(r, size=10.5, bold=(is_total or ci == 0))
        cbg(c, shade)

spc(doc, 4)

p_hdr = doc.add_paragraph()
p_hdr.paragraph_format.space_before = Pt(4)
p_hdr.paragraph_format.space_after  = Pt(2)
rh = p_hdr.add_run("\U0001f4cc  Performance Pay — How It Works:")
sf(rh, bold=True, size=10.5, color=(31, 73, 125))

perf = [
    ("Earning criteria: ",
     "Performance indicators and minimum thresholds will be jointly agreed in a detailed KPA / KPI matrix, "
     f"formalised with you and {REPORTING_TO} by the end of the probationary period."),
    ("Quarterly appraisal: ",
     "Payable at the end of each quarter, prorated for the portion of the quarter actually worked."),
    ("Carry-forward: ",
     "Any quarter where performance falls below the agreed minimum thresholds — "
     "that portion of performance pay is NOT paid and cannot be carried forward to subsequent quarters."),
    ("Post-probation: ",
     "Upon successful completion of probation, your compensation package will be reviewed "
     "and a performance pay structure aligned with company policy will be formalised."),
]
for bp, np_ in perf:
    blt(doc, bp, np_, sa=2)

spc(doc, 6)

# ════════════════════════════════════════════════
# SECTION 3 — Probationary Period
# ════════════════════════════════════════════════
h1(doc, "3.  Probationary Period")
tbl3 = doc.add_table(rows=1, cols=2)
tbl3.style = 'Table Grid'
tbl3.alignment = WD_TABLE_ALIGNMENT.LEFT
th(tbl3, ["Detail", "Information"])

prob_data = [
    ("Duration",  f"Six ({PROBATION_MONTHS}) months from date of joining"),
    ("Review",   "Performance, conduct, and overall suitability reviewed during this period"),
    ("Extension",
     f"May be extended by a further period of up to six ({PROBATION_MONTHS}) months "
     "if performance is not satisfactory, in keeping with the company's HR Policy"),
]
for idx, (k, v) in enumerate(prob_data):
    shade = 'EEF3F8' if idx % 2 == 0 else 'FFFFFF'
    row = tbl3.add_row()
    for ci, val in enumerate([k, v]):
        c = row.cells[ci]
        p = c.paragraphs[0]
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after  = Pt(3)
        r = p.add_run(val)
        sf(r, size=10.5, bold=(ci == 0))
        cbg(c, shade)

spc(doc, 6)

# ════════════════════════════════════════════════
# SECTION 4 — Office Hours, Attendance & Leave
# ════════════════════════════════════════════════
h1(doc, "4.  Office Hours, Attendance & Leave")
tbl4 = doc.add_table(rows=1, cols=2)
tbl4.style = 'Table Grid'
tbl4.alignment = WD_TABLE_ALIGNMENT.LEFT
th(tbl4, ["Category", "Terms & Conditions"])

if SATURDAY_EXEMPT:
    sat_text = (
        "Saturdays are by default non-working days.\n"
        "However, at your specific request, you have been exempted from Saturday working — "
        "subject to the conditions below:\n"
        f"   \u2022 Maximum {MAX_SATURDAYS} such Saturdays in a calendar year\n"
        "   \u2022 You may be called in or asked to work remotely for specific deliverables\n"
        "   \u2022 At least 2 working days' notice will be given wherever possible\n"
        "   \u2022 Any Saturday work will be communicated by Gowri Singh or the management"
    )
else:
    sat_text = (
        "Saturdays are 6 hours of work (Monday to Friday: 8 hours/day; "
        "Saturday: 6 hours/day, with two 30-minute breaks totalling 1 hour)."
    )

oh_data = [
    ("Office Timings",
     "Sign-in: 9:30 am – 10:00 am  |  Sign-out: 6:30 pm – 7:00 pm  (Monday to Friday)\n"
     "Total work day: 8 hours (inclusive of 1-hour total break — 30 min lunch + 30 min other)"),
    ("Late Coming",
     "Reporting later than 10:00 am for more than 3 days in a month will be treated as a half-day absent. "
     "Structured late-coming penalty policy will apply per HR Policy."),
    ("Saturdays",  sat_text),
    ("Attendance",
     "All employees sign in and sign out via the KELSA attendance system from approved locations only. "
     "Salary is computed directly from KELSA data."),
    ("Leave Entitlement",
     "\u2022 12 days — Casual / Earned Leave\n"
     "\u2022 10 days — Medical / Sick Leave\n"
     "\u2022 Public Holidays — as per the company calendar\n\n"
     "Note: Casual and Sick Leave can be availed only after the probationary period. "
     "Leave applications must be submitted at least 5 days in advance, "
     "except in cases of emergency or illness."),
]
for idx, (k, v) in enumerate(oh_data):
    row = tbl4.add_row()
    shade = 'FFF9E6' if k == "Saturdays" else ('EEF3F8' if idx % 2 == 0 else 'FFFFFF')
    for ci, val in enumerate([k, v]):
        c = row.cells[ci]
        p = c.paragraphs[0]
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after  = Pt(3)
        r = p.add_run(val)
        sf(r, size=10.5, bold=(ci == 0))
        cbg(c, shade)

spc(doc, 6)

# ════════════════════════════════════════════════
# SECTION 5 — Primary Job Responsibilities
# ════════════════════════════════════════════════
h1(doc, "5.  Primary Job Responsibilities")

body(doc,
     "You will lead the day-to-day content engine of the marketing function, working in close coordination with "
     f"{REPORTING_TO} and handing off finished assets to the Performance Marketing team for paid distribution.",
     sa=4)

subs = [
    ("A.  Content Creation & Production", [
        "Own the end-to-end content calendar — planning, scheduling, and maintaining a steady cadence of organic posts across all DRA brand channels (Instagram, Facebook, LinkedIn, YouTube, X, etc.)",
        "Ideate, design, and produce static creatives — posters, carousels, story sets, banners, brochures, and digital collaterals — in line with brand guidelines",
        "Produce short-form and long-form video content — reels, walk-throughs, project teasers, founder stories, testimonials, drone flyovers, explainers",
        "Produce animated content: stop-motion / block-motion / motion graphics / kinetic typography / 2D character animation where required",
        "Produce 3D content and collectibles using AI pipelines (2D inputs to 3D outputs). Proficiency in AI 2D-to-3D tools is required; mastery of traditional 3D modelling tools (Blender, Maya, 3ds Max) is not expected",
        "Write, edit, and review copy for social media, website, emailers, brochures, signage, and other marketing collateral",
        "Repurpose a single piece of source content into multiple platform-native formats (e.g. a 60-second reel \u2192 30-second teaser \u2192 static carousel \u2192 LinkedIn long-form post \u2192 website hero section \u2192 emailer block)",
    ]),
    ("B.  Tools, AI & Continuous Learning", [
        "Actively scout, test, and adopt new-age tools: generative AI for image/video/copy (Midjourney, Runway, Sora, Pika, Kling, ElevenLabs, GPT-class, etc.), AI 2D-to-3D tools, animation tools, and emerging platforms",
        "Working proficiency in Canva and Adobe Photoshop is expected",
        "Working knowledge of video editing tools (DaVinci Resolve, Adobe Premiere Pro, After Effects, CapCut, or equivalents) is expected — you are expected to self-learn any gap in due course",
        "Complete at least 2 pre-approved upskilling courses per year (Coursera, Udemy, Skillshare, etc.) and present learnings internally",
        "Most important trait: a HUNGRY-LEARNER mindset — willingness and speed to pick up new tools, new formats, and new platforms as the content landscape evolves",
    ]),
    ("C.  Distribution & Coordination", [
        "Publish / push out content on organic channels where appropriate, and hand off to the Performance Marketing team for paid boosting with clear briefs, captions, audience targeting, and asset variants",
        "Track performance of organic content (reach, engagement, saves, shares, follower growth) and feed learnings back into the content calendar",
    ]),
]

for sub_title, bullets in subs:
    h2(doc, sub_title)
    for b in bullets:
        blt(doc, "\u2022  ", b, sa=2)
    spc(doc, 4)

spc(doc, 4)

# ════════════════════════════════════════════════
# SECTION 6 — Secondary Responsibilities
# ════════════════════════════════════════════════
h1(doc, "6.  Secondary Responsibilities (Multi-Tasking)")

body(doc,
     "As a multi-tasking member of a lean team, you will also be expected to contribute, in coordination with "
     f"{REPORTING_TO} and the relevant function head, in the following collateral areas:",
     sa=4)

sec = [
    "Provide creative inputs and content support for interior design and architectural design collaterals — mood boards, look-books, material boards, before/after posts, lifestyle photography briefs",
    "Provide creative inputs on signage design — site hoarding, project name boards, way-finding, sales gallery, sample flat graphics — in coordination with the project / engineering team and external fabricators",
    "Identify, evaluate, and onboard content partners, agencies, freelancers, and production vendors for larger-volume requirements (drone shoots, large-format printing, model-making, voice-over, music licensing, etc.). Maintain a vetted vendor list with commercials",
    "Support the hiring process for the marketing / content function — draft JDs, source candidates, screen portfolios, and coordinate interviews with Gowri / HR",
    "Any other support function that may be reasonably required and aligned with your primary role",
]
for item in sec:
    blt(doc, "", item, sa=2)

spc(doc, 6)

# ════════════════════════════════════════════════
# SECTION 7 — KPIs
# ════════════════════════════════════════════════
h1(doc, "7.  Key Performance Indicators (KPIs)")

body(doc,
     "Your performance will be tracked against the following high-level performance areas. "
     "A detailed KPA / KPI matrix — specifying exact indicators, minimum thresholds, weightages, "
     "the formula for computing the performance-based pay component, and the cadence and format of "
     "monitoring, review, and feedback — will be formally shared with you by the end of the probationary period.",
     sa=4)

tbl7 = doc.add_table(rows=1, cols=2)
tbl7.style = 'Table Grid'
tbl7.alignment = WD_TABLE_ALIGNMENT.LEFT
th(tbl7, ["Performance Area", "Indicative Target"])

kpis = [
    ("Content Output & Cadence",  "Deliver the agreed monthly content calendar with no more than 1 missed slot per month"),
    ("Social Engagement",         "Achieve target quarterly growth in followers (combined across DRA brand profiles) and maintain a target engagement rate on posts"),
    ("Content Quality",           "Zero content requiring a public correction / clarification. Positive feedback on at least 10% of customer / audience interactions"),
    ("AI / Tool Adoption",        "Demonstrate adoption of at least 2 new tools / AI workflows per quarter, with measurable impact on output speed or quality"),
    ("Upskill",                   "Complete at least 2 approved upskilling courses per year and present learnings internally"),
    ("Cross-Function Support",   "Timely turnaround on collateral requests from interior, architectural, signage, and project teams"),
]
for idx, (k, v) in enumerate(kpis):
    shade = 'EEF3F8' if idx % 2 == 0 else 'FFFFFF'
    tr(tbl7, [k, v], shade=shade, bf=True)

spc(doc, 6)

# ════════════════════════════════════════════════
# SECTION 8 — HR Policy
# ════════════════════════════════════════════════
h1(doc, "8.  HR Policy, Code of Conduct & Confidentiality")

body(doc,
     "Your employment with DRA Realty Private Limited is governed by the company's HR Policy, including (but not limited to):",
     sa=4)

hr_items = [
    "Probation, confirmation, business hours, attendance, leave, salary advance, termination, and resignation clauses",
    "Code of Conduct (20 misconduct items) — confidentiality of company business, confidentiality of personal remuneration, prohibition on unauthorised disclosure, bribery, gambling, drunkenness, irregular attendance, willful insubordination, and strict vegetarianism in office premises",
    "Email & IT use: all emails on the company system are official records and may be reviewed or disclosed for legal / disciplinary purposes",
    "Notice period: 5 days (for employees earning < INR 20,000/month) or 30 days (for employees earning \u2265 INR 20,000/month) — applicable to both resignation and termination",
]
for item in hr_items:
    blt(doc, "\u2022  ", item, sa=2)

spc(doc, 4)

p_hr = doc.add_paragraph()
p_hr.paragraph_format.space_before = Pt(4)
p_hr.paragraph_format.space_after  = Pt(2)
rh = p_hr.add_run("\U0001f4c4  2026 DRA HR Policy:")
sf(rh, bold=True, size=10.5, color=(31, 73, 125))

p_hr2 = doc.add_paragraph()
p_hr2.paragraph_format.space_before = Pt(0)
p_hr2.paragraph_format.space_after  = Pt(4)
rh2 = p_hr2.add_run(
    "The 2026 DRA HR Policy document will be emailed to you separately (not attached to this offer) for your review and signature. "
    "Please review, sign, and return the signed copy of the HR Policy to Bharat H (sales1.blr@draas.com) as part of your "
    "joining formalities, alongside the signed copy of this offer letter.")
sf(rh2, size=10.5)

spc(doc, 6)

# ════════════════════════════════════════════════
# SECTION 9 — Required Documentation
# ════════════════════════════════════════════════
h1(doc, "9.  Required Documentation (Joining Day)")

body(doc,
     "Please bring the following original documents on your first day, along with one set of photocopies:",
     sa=4)

docs_list = [
    f"Original educational certificates ({CANDIDATE_NAME.split()[-1]}'s highest degree + B.Sc. Animation) + photocopies",
    "PAN Card (original for verification + photocopy)",
    "Aadhaar Card (photocopy)",
    "Proof of current address (utility bill / rental agreement / bank statement)",
    "Last 3 months' salary slips from prior employment (if applicable)",
    "Relieving letter and experience letter from last employer (if applicable)",
    "Two recent passport-size photographs",
    "Bank account details (cancelled cheque / passbook copy) for salary processing",
]
for item in docs_list:
    blt(doc, "", item, sa=2)

p_9b = doc.add_paragraph()
p_9b.paragraph_format.space_before = Pt(4)
p_9b.paragraph_format.space_after  = Pt(4)
r9b = p_9b.add_run(
    f"All documents may also be emailed in advance as separate, clearly-named PDF files to "
    "Bharat H (sales1.blr@draas.com), who is coordinating the onboarding process.")
sf(r9b, size=10.5, italic=True)

spc(doc, 6)

# ════════════════════════════════════════════════
# SECTION 10 — Acceptance
# ════════════════════════════════════════════════
h1(doc, "10.  Acceptance of Offer")

body(doc,
     "Kindly confirm your acceptance of this offer by replying to this email, and inform us of your proposed start date "
     "so that we may proceed with the necessary documentation and onboarding. Bharat H is being looped into this email and "
     "will take the onboarding process forward with you once we have your confirmation.\n\n"
     "Please feel free to reach out to us via email if you have any questions. Your detailed KRAs, weights, and the "
     "first-quarter content calendar will be discussed and finalised with you at the time of joining.",
     sa=6)

spc(doc, 6)
hline(doc)

p_close = doc.add_paragraph()
p_close.paragraph_format.space_before = Pt(8)
p_close.paragraph_format.space_after  = Pt(2)
rc = p_close.add_run("We are excited about the prospect of you joining the team.")
sf(rc, bold=True, size=10.5)

p_close2 = doc.add_paragraph()
p_close2.paragraph_format.space_before = Pt(0)
p_close2.paragraph_format.space_after  = Pt(16)
rc2 = p_close2.add_run("We look forward to having you on board and wish you a successful journey with us.")
sf(rc2, size=10.5)

for text, bold in [
    ("Warm regards,", False),
    ("", False),
    ("Roshini Ranka", True),
    ("Director — DRA Realty Private Limited", False),
    ("", False),
    ("Nishant Ranka", True),
    ("Director — DRA Realty Private Limited", False),
]:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(0)
    if text:
        r = p.add_run(text)
        sf(r, bold=bold, size=10.5)

spc(doc, 10)
hline(doc)

p_draft = doc.add_paragraph()
p_draft.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_draft.paragraph_format.space_before = Pt(4)
p_draft.paragraph_format.space_after  = Pt(0)
rd = p_draft.add_run(
    f"DRAFT v2 — for internal review only.  |  Candidate: {CANDIDATE_NAME} — {ROLE_TITLE}  |  "
    "Cross-ref: Prathvi Soni offer (Jan 2026) & 2026 DRA HR Policy")
sf(rd, size=8, italic=True, color=(128, 128, 128))

# ── Save ────────────────────────────────────────────────────────
import os
os.makedirs(os.path.dirname(OUTPUT_PATH), exist_ok=True)
doc.save(OUTPUT_PATH)
print(f"Saved: {OUTPUT_PATH}")