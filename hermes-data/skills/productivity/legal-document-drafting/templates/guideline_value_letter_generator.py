#!/usr/bin/env python3
"""
Template: Generate a two-letter pack for guideline value certificate + TDR confirmation
(Karnataka Sub-Registrar + JDTP/ADTP).

RUN:  python3 guideline_value_letter_generator.py
MODIFY: the constants at the bottom of this file (SURVEY_NO, VILLAGE, HOBBLI, TALUK, etc.)
  before running.

Prints: output paths of the two .docx files.
"""
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


def build_letter(ref, date, to_lines, subject, body_paras, enclosures, closing_lines, signatory, output_path):
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.2)
        section.right_margin = Inches(1)

    def add_para(text, bold=False, space_after=6, font_size=11, align=None):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(space_after)
        p.paragraph_format.space_before = Pt(0)
        if align is not None:
            p.alignment = align
        run = p.add_run(text)
        run.bold = bold
        run.font.size = Pt(font_size)
        run.font.name = "Calibri"
        return p

    add_para(ref, align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=2)
    add_para(date, align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=12)
    for line in to_lines:
        add_para(line, space_after=0)
    add_para("", space_after=4)
    add_para(f"Subject: {subject}", bold=True, space_after=12)
    add_para("Respected Sir,", space_after=10)
    for para in body_paras:
        add_para(para, space_after=10)
    if enclosures:
        add_para("Enclosures:", space_after=4)
        for e in enclosures:
            add_para(e, space_after=4)
        add_para("", space_after=4)
    for line in closing_lines:
        add_para(line, space_after=8)
    doc.add_paragraph()
    for line in signatory:
        add_para(line, bold=(line == signatory[0]))
    doc.save(output_path)
    print(f"Saved: {output_path}")


# ═══════════════════════════════════════════════════════════════════════════
#  MODIFY THESE CONSTANTS for your land parcel
# ═══════════════════════════════════════════════════════════════════════════
SURVEY_NO = "Sy. No. 14/1"          # e.g. "Sy. Nos. 14/1 to 14/7" for a range
SURVEY_NO_PLAIN = "Survey No. 14/1" # plain text version for body
VILLAGE = "Allalasandra Village"
HOBBLI = "Yelahanka Hobli"
TALUK = "Yelahanka Taluk"
DISTRICT = "Bengaluru Urban District"
LANDMARK = "immediately outside Judicial Layout, Yelahanka"  # describe location relative to known landmark
# ───────────────────────────────────────────────────────────────────────────

# L1 → Sub-Registrar
build_letter(
    ref="Ref: DRA/SR-YNK/2026-27/____",
    date="Date: ____th _______ 2026",
    to_lines=[
        "To,",
        "The Sub-Registrar,",
        "Sub-Registrar Office, Yelahanka,",
        f"{TALUK}, {DISTRICT}.",
    ],
    subject=(
        f"Request for issuance of Guideline Value Certificate / official letter in writing confirming "
        f"the guideline value (guidance value / circle rate) applicable to {SURVEY_NO} of {VILLAGE}, "
        f"{HOBBLI}, {TALUK}, {DISTRICT}"
    ),
    body_paras=[
        f"We are the owners / possessors of agricultural land comprised in {SURVEY_NO_PLAIN} of "
        f"{VILLAGE}, {HOBBLI}, {TALUK}, {DISTRICT}. The said land is situated {LANDMARK} and forms part "
        f"of {VILLAGE} proper. The attested RTC copy of the above survey number is enclosed for ready reference.",
        "We propose to proceed with development of the said land and, in connection therewith, we are required to load "
        "Transferable Development Rights (TDR) on the same. For this purpose, as also for computation of the applicable "
        "stamp duty and registration charges, the officially notified guideline value applicable to the said survey number is required.",
        "We, therefore, request you to kindly issue an official letter / certificate in writing confirming the following:",
        "a)  the guideline value (guidance value / circle rate) per sq. m. / per sq. ft. notified and applicable to "
        f"{SURVEY_NO} of {VILLAGE} as per the current notification issued under the Karnataka Stamp Act, 1957;",
        "b)  the classification of the land (agricultural / non-agricultural) and the guidance-value zone in which the land falls;",
        "c)  the reference number and date of the notification under which the said guideline value is prescribed.",
        "We further request you to confirm that the value of the TDR to be loaded on the said land will be computed on the basis of "
        "the aforesaid guideline value.",
        "We shall be grateful for your kind consideration and issuance of the said certificate at the earliest, so that we may "
        "proceed with the necessary registrations and statutory payments.",
    ],
    enclosures=[
        f"1. Attested RTC copy of {SURVEY_NO}, {VILLAGE}, {HOBBLI}, {TALUK}, {DISTRICT}.",
        "2. [Village map / location sketch, if any]",
    ],
    closing_lines=["Thanking you,", "Yours faithfully,"],
    signatory=[
        "______________________",
        "[Name of Applicant / Authorised Signatory]",
        "[Designation / Company]",
        "Contact: [Mobile No.] / [Email]",
    ],
    output_path=f"/tmp/L1_Guideline_Value_SubRegistrar_{SURVEY_NO.replace(' ','_').replace('/','-')}.docx",
)

# L2 → JDTP/ADTP
build_letter(
    ref="Ref: DRA/TP/2026-27/____",
    date="Date: ____th _______ 2026",
    to_lines=[
        "To,",
        "The Joint Director of Town Planning (JDTP) /",
        "Assistant Director of Town Planning (ADTP),",
        "Office of Town Planning, Bengaluru Urban District,",
        "Karnataka.",
        "Attn: [Shri Nagrajappa]",
    ],
    subject=(
        f"Confirmation of guideline value and basis of computation of TDR value for land bearing "
        f"{SURVEY_NO} of {VILLAGE}, {HOBBLI}, {TALUK}, {DISTRICT}"
    ),
    body_paras=[
        f"This has reference to the land belonging to us comprising {SURVEY_NO_PLAIN} of "
        f"{VILLAGE}, {HOBBLI}, {TALUK}, {DISTRICT}, situated {LANDMARK}. The attested RTC copy of the "
        "said survey number is enclosed for ready reference.",
        "We have separately applied to the Sub-Registrar, Yelahanka for issuance of the official letter / certificate "
        "confirming the guideline value applicable to the said land.",
        "In this connection, we request your office to kindly confirm in writing:",
        f"a)  the guideline value applicable to {SURVEY_NO} as per the current notification under the Karnataka Stamp Act, "
        "1957, and that the same is the confirmed guideline value being adopted by the authorities;",
        "b)  that the value of the Transferable Development Rights (TDR) to be loaded on the said land will be computed "
        "on the basis of the said guideline value; and",
        "c)  the notification reference and the method / formula applied for computation of the TDR value.",
        "We shall be grateful for your written confirmation at the earliest, as the same is required for our statutory "
        "filings and development approvals.",
    ],
    enclosures=[
        f"1. Attested RTC copy of {SURVEY_NO}, {VILLAGE}, {HOBBLI}, {TALUK}, {DISTRICT}.",
    ],
    closing_lines=["Thanking you,", "Yours faithfully,"],
    signatory=[
        "______________________",
        "[Name of Applicant / Authorised Signatory]",
        "[Designation / Company]",
        "Contact: [Mobile No.] / [Email]",
    ],
    output_path=f"/tmp/L2_TDR_Guideline_Value_Confirmation_JDTP_ADTP_{SURVEY_NO.replace(' ','_').replace('/','-')}.docx",
)