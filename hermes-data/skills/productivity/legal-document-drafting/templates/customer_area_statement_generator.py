#!/usr/bin/env python3
"""
DRA Realty Pvt. Ltd. letterhead — Customer Area Statement (bank project pre-approval).
Reads raw sheet data from /tmp/area_statement_raw.json.
Portrait covering letter + landscape Customer Area Statement annexure (Annexure-A),
structured per the Customer Area Statement spec: RERA Carpet -> Exclusive Areas
-> BUA (walls incl.) -> Common Loading (+%) -> Super BUA -> UDS.

USAGE:  python3 this_script.py  # outputs to /tmp/...
"""
import json
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NAVY = RGBColor(0x1F, 0x38, 0x64)
GOLD = RGBColor(0xC9, 0x9A, 0x2E)
GREY = RGBColor(0x55, 0x55, 0x55)
LOGO = "/data/hermes/cache/analysis/dra_logo/dra-logo.png"

with open('/tmp/area_statement_raw.json') as f:
    ROWS = json.load(f)

DATA = ROWS[1:21]   # 20 units
TOTAL = ROWS[21]    # totals row

def num(v):
    return float(str(v).replace(',', '').replace('%', '').strip() or 0)

def loading_pct(row):
    bua = num(row[7]); sba = num(row[13])
    if bua <= 0:
        return 0.0
    return (sba - bua) / bua * 100.0

def set_cell_bg(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:fill'), hexcolor)
    tcPr.append(shd)

def set_cell_borders(cell, color="BFBFBF", sz="4"):
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        el = OxmlElement(f'w:{edge}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), sz)
        el.set(qn('w:color'), color)
        borders.append(el)
    tcPr.append(borders)

def cell_text(cell, text, bold=False, size=8, color=None, align='left', font='Calibri'):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    if align == 'center':
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == 'right':
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = font
    if color is not None:
        run.font.color.rgb = color

def add_rule(paragraph, color="1F3864", sz="18"):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), sz)
    bottom.set(qn('w:space'), '4')
    bottom.set(qn('w:color'), color)
    pBdr.append(bottom)
    pPr.append(pBdr)

doc = Document()
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(10.5)

# ════════════════════════════════════════════════════════════════════════════
# SECTION 1 — PORTRAIT LETTERHEAD
# ════════════════════════════════════════════════════════════════════════════
sec1 = doc.sections[0]
sec1.page_width = Cm(21.0)
sec1.page_height = Cm(29.7)
sec1.top_margin = Inches(0.55)
sec1.bottom_margin = Inches(0.6)
sec1.left_margin = Inches(0.8)
sec1.right_margin = Inches(0.8)

tbl = doc.add_table(rows=1, cols=2)
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
tbl.autofit = False
tbl.columns[0].width = Cm(5.2)
tbl.columns[1].width = Cm(13.6)
lc, rc = tbl.rows[0].cells
lc.width = Cm(5.2)
rc.width = Cm(13.6)

p = lc.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
run = p.add_run()
run.add_picture(LOGO, width=Cm(4.6))

rc.paragraphs[0].text = ""
p1 = rc.paragraphs[0]
p1.alignment = WD_ALIGN_PARAGRAPH.RIGHT
r = p1.add_run("DRA REALTY PVT. LTD.")
r.bold = True
r.font.size = Pt(20)
r.font.color.rgb = NAVY
p2 = rc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
r = p2.add_run("\u201CHOME OF PRIDE\u201D")
r.font.size = Pt(11)
r.font.color.rgb = GOLD
r.bold = True
p3 = rc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.RIGHT
r = p3.add_run("201A / 202BA Queens Corner, No. 3, Queens Road, Bengaluru \u2013 560 001\n")
r.font.size = Pt(8.5)
r.font.color.rgb = GREY
r2 = p3.add_run("Tel: +91 98800 55634  |  Email: ndr@draas.com  |  www.drahomes.in")
r2.font.size = Pt(8.5)
r2.font.color.rgb = GREY

rule = doc.add_paragraph()
rule.paragraph_format.space_after = Pt(10)
add_rule(rule, color="1F3864", sz="18")
rule2 = doc.add_paragraph()
rule2.paragraph_format.space_after = Pt(16)
rule2.paragraph_format.space_before = Pt(0)
add_rule(rule2, color="C99A2E", sz="8")

pref = doc.add_paragraph()
pref.alignment = WD_ALIGN_PARAGRAPH.RIGHT
pref.paragraph_format.space_after = Pt(2)
r = pref.add_run("Ref: DRA/BANK/PRE-APPROVAL/____")
r.font.size = Pt(10)
pdt = doc.add_paragraph()
pdt.alignment = WD_ALIGN_PARAGRAPH.RIGHT
pdt.paragraph_format.space_after = Pt(16)
r = pdt.add_run("Date: ____th ________ 2026")
r.font.size = Pt(10)

to_lines = [
    "To,",
    "The Manager,",
    "[Name of Bank / Branch],",
    "[Branch Address],",
    "Bengaluru.",
]
for line in to_lines:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    r = p.add_run(line)
    r.font.size = Pt(10.5)

psub = doc.add_paragraph()
psub.paragraph_format.space_before = Pt(14)
psub.paragraph_format.space_after = Pt(12)
r = psub.add_run("Subject: Submission of Customer Area Statement \u2013 [Project Name] \u2013 Request for Project Pre-Approval")
r.bold = True
r.font.size = Pt(11.5)
r.font.color.rgb = NAVY

doc.add_paragraph().paragraph_format.space_after = Pt(2)
psal = doc.add_paragraph()
psal.paragraph_format.space_after = Pt(10)
r = psal.add_run("Dear Sir,")
r.font.size = Pt(10.5)

body = [
    "We are pleased to submit the Customer Area Statement of our residential project \u201c[Project Name]\u201d, "
    "located at [Address], Bengaluru, for your kind perusal and "
    "consideration of project finance / pre-approval.",
    "The project comprises [N] residential units ([Configuration]), with a total saleable "
    "area of [X] sq.ft as per the sanctioned plan and the Supplementary Sharing Agreement, on a total "
    "sanctioned plot area of [Y] sq.ft. The unit-wise area break-up, including RERA Carpet Area, "
    "exclusive balcony areas, Built-up Area, common area loading (with loading percentage), Super Built-Up "
    "Area and undivided share of land (UDS), is provided in the enclosed Customer Area Statement (Annexure-A).",
    "We request you to kindly process the same at the earliest and revert with your requirements, if any, "
    "for the project pre-approval. We shall be happy to provide any further documents or clarifications "
    "that may be required.",
]
for para in body:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.line_spacing = 1.15
    r = p.add_run(para)
    r.font.size = Pt(10.5)

pe = doc.add_paragraph()
pe.paragraph_format.space_after = Pt(2)
r = pe.add_run("Enclosure: ")
r.bold = True
r.font.size = Pt(10.5)
r2 = pe.add_run("Annexure-A \u2013 Customer Area Statement ([N] units, [Project Name])")
r2.font.size = Pt(10.5)

for line in ["Thanking you,", "Yours faithfully,", "For DRA Realty Pvt. Ltd."]:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(line)
    r.font.size = Pt(10.5)

for _ in range(2):
    doc.add_paragraph()
sig_lines = [
    ("________________________", False),
    ("Nishant Ranka", True),
    ("Managing Director", False),
    ("DRA Realty Pvt. Ltd.", False),
]
for text, bold in sig_lines:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text)
    r.font.size = Pt(10.5)
    r.bold = bold

# ════════════════════════════════════════════════════════════════════════════
# SECTION 2 — LANDSCAPE CUSTOMER AREA STATEMENT (ANNEXURE-A)
# ════════════════════════════════════════════════════════════════════════════
sec2 = doc.add_section(WD_SECTION.NEW_PAGE)
sec2.orientation = WD_ORIENT.LANDSCAPE
sec2.page_width = Cm(29.7)
sec2.page_height = Cm(21.0)
sec2.top_margin = Inches(0.5)
sec2.bottom_margin = Inches(0.5)
sec2.left_margin = Inches(0.5)
sec2.right_margin = Inches(0.5)

pt = doc.add_paragraph()
pt.alignment = WD_ALIGN_PARAGRAPH.CENTER
pt.paragraph_format.space_after = Pt(2)
r = pt.add_run("ANNEXURE-A")
r.bold = True
r.font.size = Pt(12)
r.font.color.rgb = NAVY
pt2 = doc.add_paragraph()
pt2.alignment = WD_ALIGN_PARAGRAPH.CENTER
pt2.paragraph_format.space_after = Pt(4)
r = pt2.add_run("CUSTOMER AREA STATEMENT \u2013 [PROJECT NAME]")
r.bold = True
r.font.size = Pt(13)
r.font.color.rgb = NAVY

pm = doc.add_paragraph()
pm.alignment = WD_ALIGN_PARAGRAPH.CENTER
pm.paragraph_format.space_after = Pt(6)
r = pm.add_run("Project: [Project Name] | Location: [Site Address] | "
               "Configuration: [Configuration] | Total Units: [N] | Plot Area: [Y] sq.ft.")
r.font.size = Pt(9)
r.font.color.rgb = GREY

# Formula block
pf = doc.add_paragraph()
pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
pf.paragraph_format.space_after = Pt(8)
r = pf.add_run("Super Built-up Area = RERA Carpet Area + Exclusive Balcony/Utility Area + External Walls + Common Area Loading")
r.bold = True
r.font.size = Pt(9)
r.font.color.rgb = NAVY

# Table: 12 cols (adjust col_widths as needed for your columns)
ncols = 12
table = doc.add_table(rows=1 + len(DATA) + 1, cols=ncols)
table.style = 'Table Grid'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.autofit = False

col_widths = [Cm(0.7), Cm(1.0), Cm(0.8), Cm(1.4), Cm(1.5), Cm(1.8), Cm(1.7), Cm(1.7), Cm(1.7), Cm(1.4), Cm(1.7), Cm(1.8)]

hdrs = ["#", "Unit #", "Share", "Floor", "Entrance\nFacing (Boundary)", "RERA Carpet\nArea (sft)", "Balcony / Exclusive\nArea (sft)", "Built-up Area\n(sft) [Walls incl.]", "Common Area\nLoading (sft)", "Loading\n(%)", "Super Built-up\nArea (sft)", "UDS\n(sft)"]

for j, h in enumerate(hdrs):
    cell = table.cell(0, j)
    cell.width = col_widths[j]
    set_cell_bg(cell, "1F3864")
    set_cell_borders(cell, color="1F3864")
    cell_text(cell, h, bold=True, size=7.5, color=RGBColor(0xFF, 0xFF, 0xFF), align='center')

for i, row in enumerate(DATA, start=1):
    vals = [
        str(row[0]).strip(),
        str(row[1]).strip(),
        str(row[2]).strip(),
        str(row[4]).strip(),
        str(row[5]).strip(),
        f"{num(row[10]):,.0f}",
        f"{num(row[9]):,.0f}",
        f"{num(row[7]):,.0f}",
        f"{num(row[12]):,.0f}",
        f"{loading_pct(row):.1f}%",
        f"{num(row[13]):,.0f}",
        f"{num(row[14]):,.2f}",
    ]
    for j, v in enumerate(vals):
        cell = table.cell(i, j)
        cell.width = col_widths[j]
        set_cell_borders(cell)
        if i % 2 == 0:
            set_cell_bg(cell, "F4F1E8")
        align = 'right' if j >= 5 else 'center'
        cell_text(cell, v, size=8, align=align)

# totals row
trow = 1 + len(DATA)
tot_vals = [
    "", "", "", "", "",
    f"{sum(num(r[10]) for r in DATA):,.0f}",
    f"{sum(num(r[9]) for r in DATA):,.0f}",
    f"{sum(num(r[7]) for r in DATA):,.0f}",
    f"{sum(num(r[12]) for r in DATA):,.0f}",
    "",
    f"{num(TOTAL[13]):,.0f}",
    f"{num(TOTAL[14]):,.2f}",
]
for j, v in enumerate(tot_vals):
    cell = table.cell(trow, j)
    cell.width = col_widths[j]
    set_cell_bg(cell, "DCE3F0")
    set_cell_borders(cell)
    cell_text(cell, "Total" if j == 3 else v, bold=True, size=8, align='right' if j >= 5 else 'center')

# Notes
fn = doc.add_paragraph()
fn.paragraph_format.space_before = Pt(8)
fn.paragraph_format.space_after = Pt(2)
r = fn.add_run("Notes: ")
r.bold = True
r.font.size = Pt(8.5)
r2 = fn.add_run("[customise: saleable area, super BUA, UDS/sft, share definitions]")
r2.font.size = Pt(8.5)
r2.font.color.rgb = GREY

fn2 = doc.add_paragraph()
fn2.paragraph_format.space_after = Pt(2)
r = fn2.add_run("External Walls & Exclusions: ")
r.bold = True
r.font.size = Pt(8.5)
r2 = fn2.add_run("External wall thickness and structural columns are included within the Built-up Area figures above as per the "
                 "sanctioned plan. Utility / wash areas and exclusive open terrace: [add if applicable]. "
                 "Car parking charges, clubhouse membership fees, GST, registration costs and maintenance deposits "
                 "are NOT included in the carpet area / apartment price.")
r2.font.size = Pt(8.5)
r2.font.color.rgb = GREY

fn3 = doc.add_paragraph()
r = fn3.add_run("This Customer Area Statement is furnished in connection with the project pre-approval request and is subject to "
                "verification against the sanctioned plan and the definitive project documents.")
r.font.size = Pt(8.5)
r.font.color.rgb = GREY

out = '/tmp/DRA_Realty_[Project]_Customer_Area_Statement_Bank_Letter.docx'
doc.save(out)
print("Saved:", out)