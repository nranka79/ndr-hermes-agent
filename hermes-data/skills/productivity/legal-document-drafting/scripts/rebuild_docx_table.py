"""
rebuild_docx_table.py
Add or replace a column in an existing python-docx table via raw XML manipulation.

Use when: python-docx lacks insert_column(); you need to add a column to an existing table
(e.g., Super Built-up Area from cost sheet added to Schedule B in Ranka Amber, May 2026).

Pattern (5 steps):
  1. Read existing data: [c.paragraphs[0].text for c in row.cells] per row
  2. Remove all <w:tr> children from <w:tbl> via lxml/etree
  3. Clear and rebuild <w:tblGrid> with new column count and widths
  4. Append fresh <w:tr> rows with all columns via make_cell() XML construction
  5. doc.save() — verify with new Document(path); print(len(t2.columns))

Dependencies:
  from docx import Document
  from docx.oxml.ns import qn
  from docx.oxml import OxmlElement
  from lxml import etree
  from copy import deepcopy

Column width (dxa = twentieths of a point):
  900 dxa ≈ 1 cm
  1814 dxa ≈ 2 cm (original schedule column width)
  2100 dxa ≈ 2.3 cm (good for wider columns like Super BUA sq.ft)

Color fill constants:
  HDR_FILL = 'B4C6E7'   # blue header row
  DAT_FILL = 'FFF2CC'   # yellow data row

Live use case (Ranka Amber Schedule B — May 2026):
  cost_super_bu = {'101': 1687, '102': 1712, ...}
  # Table 1 = Schedule B, 6 cols → 7 cols
  # Table 2 = Schedule C (Landowner), 5 cols → 6 cols
  # Table 3 = Schedule D (Developer), 5 cols → 6 cols
"""

from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from lxml import etree
from copy import deepcopy

# ──────────────────────────────────────────────────────────────
# Helper functions
# ──────────────────────────────────────────────────────────────

def make_para(text, bold=False, align='left'):
    p = OxmlElement('w:p')
    pPr = OxmlElement('w:pPr')
    jc = OxmlElement('w:jc')
    jc.set(qn('w:val'), align)
    pPr.append(jc)
    p.append(pPr)
    r = OxmlElement('w:r')
    if bold:
        rPr = OxmlElement('w:rPr')
        b = OxmlElement('w:b')
        rPr.append(b)
        r.append(rPr)
    t = OxmlElement('w:t')
    t.text = str(text)
    t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    r.append(t)
    p.append(r)
    return p

def make_cell(text, w, fill, align, bold=None):
    if bold is None:
        bold = (fill == 'B4C6E7')  # HDR_FILL
    tc = OxmlElement('w:tc')
    tcPr = OxmlElement('w:tcPr')
    tcW = OxmlElement('w:tcW')
    tcW.set(qn('w:w'), str(w))
    tcW.set(qn('w:type'), 'dxa')
    tcPr.append(tcW)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:fill'), fill)
    tcPr.append(shd)
    tc.append(tcPr)
    tc.append(make_para(text, bold=bold, align=align))
    return tc

def make_header_row(widths, headers, hdr_fill='B4C6E7', hdr_align='center'):
    tr = OxmlElement('w:tr')
    for h, w in zip(headers, widths):
        tc = make_cell(h, w, hdr_fill, hdr_align, bold=True)
        tr.append(tc)
    return tr

def make_data_row(cells_data, widths, dat_fill='FFF2CC', dat_align='center'):
    tr = OxmlElement('w:tr')
    for cell_text, w in zip(cells_data, widths):
        tc = make_cell(cell_text, w, dat_fill, dat_align, bold=False)
        tr.append(tc)
    return tr

def rebuild_table_with_new_column(table, col_widths, headers, fill_func=None, data_rows=None):
    """
    Rebuild an existing table with a new column structure.

    Args:
        table: doc.tables[idx]
        col_widths: list of int dxa values (one per column)
        headers: list of str header texts
        fill_func: callable(unit_no) -> str  — provides data for new column per row
                  If None, data_rows must be provided.
        data_rows: list of list of str — explicit row data (all columns)
                  Only used if fill_func is None.
    """
    tbl = table._tbl

    # Collect existing row data
    rows_data = []
    for row in table.rows:
        cells = [c.paragraphs[0].text for c in row.cells]
        rows_data.append(cells)

    # Remove all existing rows
    for child in list(tbl):
        if etree.QName(child).localname == 'tr':
            tbl.remove(child)

    # Update tblGrid
    tblGrid = tbl.find(qn('w:tblGrid'))
    for gc in list(tblGrid):
        tblGrid.remove(gc)
    for w in col_widths:
        gc = OxmlElement('w:gridCol')
        gc.set(qn('w:w'), str(w))
        tblGrid.append(gc)

    # Header row
    tbl.append(make_header_row(col_widths, headers))

    # Data rows
    for row_data in rows_data[1:]:  # skip header
        if fill_func:
            unit_no = row_data[0]  # assumes unit is first column
            new_col_val = fill_func(unit_no)
            cells_data = row_data + [new_col_val]
        else:
            cells_data = data_rows[len(rows_data) - 1] if data_rows else row_data
        tbl.append(make_data_row(cells_data, col_widths))

    # Update table width
    total_w = sum(col_widths)
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is not None:
        tblW = tblPr.find(qn('w:tblW'))
        if tblW is not None:
            tblW.set(qn('w:w'), str(total_w))
            tblW.set(qn('w:type'), 'dxa')


# ──────────────────────────────────────────────────────────────
# Ranka Amber — Add Super BUA to Schedule B, C, D (May 2026)
# ──────────────────────────────────────────────────────────────

cost_super_bu = {
    '101': 1687, '102': 1712, '103': 1468, '104': 1219, '105': 1504,
    '201': 1792, '202': 1830, '203': 1572, '204': 1276, '205': 1617,
    '301': 1792, '302': 1830, '303': 1572, '304': 1276, '305': 1617,
    '401': 1792, '402': 1830, '403': 1572, '404': 1276, '405': 1617,
}

allotted_map = {
    '101': 'Landowner (LO)', '102': 'Landowner (LO)', '103': 'Landowner (LO)',
    '104': 'Landowner (LO)', '105': 'Landowner (LO)',
    '201': 'Developer (DEV)', '202': 'Developer (DEV)', '203': 'Developer (DEV)',
    '204': 'Developer (DEV)', '205': 'Developer (DEV)',
    '301': 'Developer (DEV)', '302': 'Developer (DEV)', '303': 'Developer (DEV)',
    '304': 'Developer (DEV)', '305': 'Developer (DEV)',
    '401': 'Landowner (LO)', '402': 'Landowner (LO)', '403': 'Landowner (LO)',
    '404': 'Landowner (LO)', '405': 'Landowner (LO)',
}

if __name__ == '__main__':
    import sys
    doc_path = sys.argv[1] if len(sys.argv) > 1 else '/data/hermes/cron/output/Ranka_Amber_Supplementary_Sharing_Agreement_Draft.docx'
    doc = Document(doc_path)

    # ── Schedule B (Table 1) — 6 existing cols + Allotted To = 7 cols ──
    # Headers: Floor | Plan Unit No. | Marketing Unit No. | BUA(sq.m) | Carpet(sq.m) | Super BUA(sq.ft) | Allotted To
    tB = doc.tables[1]
    tblB = tB._tbl
    rows_B = [[c.paragraphs[0].text for c in r.cells] for r in tB.rows]
    for child in list(tblB):
        if etree.QName(child).localname == 'tr':
            tblB.remove(child)
    tblGridB = tblB.find(qn('w:tblGrid'))
    for gc in list(tblGridB): tblGridB.remove(gc)
    col_w_B = [900, 1000, 1600, 1500, 1500, 1700, 1500]
    for w in col_w_B:
        gc = OxmlElement('w:gridCol'); gc.set(qn('w:w'), str(w)); tblGridB.append(gc)
    hdr_B = ['Floor', 'Plan Unit No.', 'Marketing Unit No.', 'Built-up Area (sq.m)',
             'Carpet Area (sq.m)', 'Super Built-up Area (sq.ft)', 'Allotted To']
    tblB.append(make_header_row(col_w_B, hdr_B))
    for rd in rows_B[1:]:
        unit = rd[2]
        tblB.append(make_data_row(rd[:5] + [str(cost_super_bu.get(unit,'')), allotted_map.get(unit,'')], col_w_B))

    # ── Schedule C (Table 2) — 5 cols + Super BUA = 6 cols ──
    tC = doc.tables[2]
    tblC = tC._tbl
    rows_C = [[c.paragraphs[0].text for c in r.cells] for r in tC.rows]
    for child in list(tblC):
        if etree.QName(child).localname == 'tr':
            tblC.remove(child)
    tblGridC = tblC.find(qn('w:tblGrid'))
    for gc in list(tblGridC): tblGridC.remove(gc)
    col_w_C = [1500, 1500, 1500, 1500, 1500, 2100]
    for w in col_w_C:
        gc = OxmlElement('w:gridCol'); gc.set(qn('w:w'), str(w)); tblGridC.append(gc)
    hdr_C = ['Marketing Unit No.', 'Floor', 'Plan Unit No.',
             'Built-up Area (sq.m)', 'Carpet Area (sq.m)', 'Super Built-up Area (sq.ft)']
    tblC.append(make_header_row(col_w_C, hdr_C))
    for rd in rows_C[1:]:
        unit = rd[0]
        tblC.append(make_data_row(rd + [str(cost_super_bu.get(unit,''))], col_w_C))

    # ── Schedule D (Table 3) — 5 cols + Super BUA = 6 cols ──
    tD = doc.tables[3]
    tblD = tD._tbl
    rows_D = [[c.paragraphs[0].text for c in r.cells] for r in tD.rows]
    for child in list(tblD):
        if etree.QName(child).localname == 'tr':
            tblD.remove(child)
    tblGridD = tblD.find(qn('w:tblGrid'))
    for gc in list(tblGridD): tblGridD.remove(gc)
    col_w_D = [1500, 1500, 1500, 1500, 1500, 2100]
    for w in col_w_D:
        gc = OxmlElement('w:gridCol'); gc.set(qn('w:w'), str(w)); tblGridD.append(gc)
    hdr_D = ['Marketing Unit No.', 'Floor', 'Plan Unit No.',
             'Built-up Area (sq.m)', 'Carpet Area (sq.m)', 'Super Built-up Area (sq.ft)']
    tblD.append(make_header_row(col_w_D, hdr_D))
    for rd in rows_D[1:]:
        unit = rd[0]
        tblD.append(make_data_row(rd + [str(cost_super_bu.get(unit,''))], col_w_D))

    doc.save(doc_path)
    print('Saved. Verifying...')
    doc2 = Document(doc_path)
    for idx, t in enumerate([doc2.tables[1], doc2.tables[2], doc2.tables[3]]):
        hdr = [c.paragraphs[0].text for c in t.rows[0].cells]
        row1 = [c.paragraphs[0].text for c in t.rows[1].cells]
        print('Table %d: cols=%d header=%s row1=%s' % (idx, len(t.columns), hdr, row1))