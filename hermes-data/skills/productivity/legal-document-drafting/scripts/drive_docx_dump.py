#!/usr/bin/env python3
"""
Dump a Google-Drive-hosted .docx (.docx upload, not a native Google Doc) to plain text.

WHY: Google Docs API refuses to parse native Office files even when the user shares a
docs.google.com/document/d/<id> link. Symptom:
  HttpError 400 "This operation is not supported for this document. The document must not be an Office file."
and Drive export also fails: "Export only supports Docs Editors files."
The file is a binary .docx stored on Drive; the working recipe is:

  1. drive.files().get(fileId=...)  -> confirm mimeType == application/vnd.openxmlformats-officedocument.wordprocessingml.document
  2. drive.files().get_media(fileId=...) -> download raw bytes
  3. parse with python-docx from /opt/data/.venv-docx/bin/python3 (verified working, has docx + lxml)

Usage:
  /opt/hermes/.venv/bin/python3 drive_docx_dump.py <FILE_ID> [output.txt]

Prints paragraph text (as [P<n>]) then all tables (rows x cols) to stdout,
and writes the same to the optional output file.
"""
import sys
sys.path.insert(0, "/opt/hermes")

file_id = sys.argv[1] if len(sys.argv) > 1 else None
outfile = sys.argv[2] if len(sys.argv) > 2 else None
if not file_id:
    print("usage: drive_docx_dump.py <FILE_ID> [output.txt]")
    sys.exit(1)

from tools.gws_auth import build_service

drive = build_service("drive", "v3", service_name="google-draas")
meta = drive.files().get(fileId=file_id, fields="id, name, mimeType, size, parents, modifiedTime").execute()
print(f"FILE: {meta['name']} | mime: {meta['mimeType']} | size: {meta.get('size')}")
if meta["mimeType"] != "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
    print("NOTE: not a binary .docx — try Docs API instead if it's a native Google Doc")

raw = drive.files().get_media(fileId=file_id).execute()
tmp_docx = "/tmp/_drive_docx_source.docx"
with open(tmp_docx, "wb") as f:
    f.write(raw)

# Parse with the docx venv (has python-docx + lxml). The /opt/hermes venv often lacks python-docx.
import subprocess
code = r'''
import docx, sys
d = docx.Document("/tmp/_drive_docx_source.docx")
out = []
for i, p in enumerate(d.paragraphs):
    t = p.text.strip()
    if t:
        out.append(f"[P{i}] {t}")
for ti, tbl in enumerate(d.tables):
    out.append(f"--- TABLE {ti} ({len(tbl.rows)}x{len(tbl.columns)}) ---")
    for row in tbl.rows:
        out.append(" | ".join(c.text.strip().replace("\n", " / ") for c in row.cells))
print("\n".join(out))
'''
res = subprocess.run(["/opt/data/.venv-docx/bin/python3", "-c", code], capture_output=True, text=True, timeout=180)
if res.returncode != 0:
    print("PARSE ERROR:", res.stderr[-3000:])
    sys.exit(1)
text = res.stdout
print(text)
if outfile:
    with open(outfile, "w") as f:
        f.write(text)
    print(f"\n[written to {outfile}]")