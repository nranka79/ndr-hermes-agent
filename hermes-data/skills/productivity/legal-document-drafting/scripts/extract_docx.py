#!/usr/bin/env python3
"""Extract full text content from a .docx (paragraphs + tables).

Usage: python extract_docx.py <file.docx>
Requires python-docx (pip install python-docx).

Use when you need to inspect a legal document's structure before editing:
paragraph indices, table positions, placeholder locations.
"""
import sys

from docx import Document


def main(path):
    doc = Document(path)
    print(f"# PARAGRAPHS: {len(doc.paragraphs)}")
    for i, p in enumerate(doc.paragraphs):
        t = p.text.strip()
        if t:
            print(f"[{i}] {t}")
    print(f"# TABLES: {len(doc.tables)}")
    for ti, table in enumerate(doc.tables):
        print(f"--- TABLE {ti} ({len(table.rows)} rows x {len(table.columns)} cols) ---")
        for ri, row in enumerate(table.rows):
            print(f"{ri}: {[c.text.strip().replace(chr(10), ' / ') for c in row.cells]}")


if __name__ == "__main__":
    main(sys.argv[1])