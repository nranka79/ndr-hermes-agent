#!/usr/bin/env python3
"""
BBMP Letter Template — Ranka Iris License Fee Letter
Pattern: right-aligned Ref/Date → To block → bold Subject → Salutation → body → optional annexure table → closing → signature

Usage: copy and modify text values, then run:
  python3 /data/hermes/skills/productivity/legal-document-drafting/scripts/bbmp_letter_template.py
"""
from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_bbmp_letter(ref, date, to_lines, subject, body_paras, annexures, closing_lines, signatory, output_path="/tmp/letter.docx"):
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.2)
        section.right_margin = Inches(1)

    def add_para(text, bold=False, space_after=6, font_size=11):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(space_after)
        run = p.add_run(text)
        run.bold = bold
        run.font.size = Pt(font_size)
        run.font.name = "Calibri"
        return p

    def add_line(text, bold=False, font_size=11):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        run = p.add_run(text)
        run.bold = bold
        run.font.size = Pt(font_size)
        run.font.name = "Calibri"
        return p

    # Ref (right)
    p_ref = doc.add_paragraph()
    p_ref.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_ref.paragraph_format.space_after = Pt(2)
    r = p_ref.add_run(ref)
    r.font.size = Pt(11)
    r.font.name = "Calibri"

    # Date (right)
    p_date = doc.add_paragraph()
    p_date.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_date.paragraph_format.space_after = Pt(12)
    r = p_date.add_run(date)
    r.font.size = Pt(11)
    r.font.name = "Calibri"

    # To block
    for line in to_lines:
        add_line(line)

    # Subject
    p_sub = doc.add_paragraph()
    p_sub.paragraph_format.space_before = Pt(12)
    p_sub.paragraph_format.space_after = Pt(12)
    r = p_sub.add_run(f"Subject: {subject}")
    r.bold = True
    r.font.size = Pt(11)
    r.font.name = "Calibri"

    # Salutation
    add_para("Respected Sir,", space_after=10)

    # Body paragraphs
    for para in body_paras:
        add_para(para, space_after=10)

    # Optional annexure table
    if annexures:
        table = doc.add_table(rows=1, cols=2)
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        hdr[0].text = "Annexure"
        hdr[1].text = "Description"
        for cell in hdr:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.bold = True
                    run.font.size = Pt(10)
                    run.font.name = "Calibri"
        for label, desc in annexures:
            row = table.add_row()
            row.cells[0].text = label
            row.cells[1].text = desc
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.size = Pt(10)
                        run.font.name = "Calibri"
        table.columns[0].width = Cm(3)
        table.columns[1].width = Cm(13)
        doc.add_paragraph()

    # Closing
    for line in closing_lines:
        add_para(line, space_after=8)

    doc.add_paragraph()

    # Signature block
    for line in signatory:
        add_line(line, bold=(line == signatory[0]))

    doc.save(output_path)
    print(f"Saved: {output_path}")


# ─── CALL WITH ACTUAL VALUES ─────────────────────────────────────────────────
if __name__ == "__main__":
    create_bbmp_letter(
        ref="Ref: DRA/BBMP/2026-27/___",
        date="Date: 25th May 2026",
        to_lines=[
            "To,",
            "The Joint Director (Town Planning – North),",
            "Bruhat Bengaluru Mahanagara Palike (BBMP),",
            "Annex-3 Building, 4th Floor, N.R. Square,",
            "Bengaluru – 560 002.",
        ],
        subject="Respectful submission regarding License Fee for Ranka Iris Project – BBMP/Sanction Plan No. BBMP/Addl.Dir/JD NORTH/0037/2013-14.",
        body_paras=[
            "Warm greetings from DRA Developers & Projects Pvt. Ltd. We sincerely thank you and the esteemed BBMP Town Planning department for processing our Occupancy Certificate application.",
            "We are in receipt of the official Occupancy Certificate Fee Demand Challan (Document No. BBMP/Addl.Dir/JD North/LP/0037/2013-14, dated 30-04-2026, attached as Annexure A).",
            "With respect to the License Fee component included in the demand, we wish to submit that the License Fee for the first two terms has already been remitted by us, as per the details below. We request your kind consideration of the same while processing the fee challan:",
        ],
        annexures=[
            ("Annexure A", "Occupancy Certificate Fee Demand Challan (Doc No. BBMP/Addl.Dir/JD North/LP/0037/2013-14, Dated 30-04-2026)"),
            ("Annexure B", "Proof of earlier License Fee payments:\n• Term 1 (Sept 2013 – Sept 2018): Paid at Plan Sanction Stage via DD No. 185377 dated 08-08-2013.\n• Term 2 (Sept 2018 – Sept 2023): Paid at Commencement Certificate Stage via DD No. 506078 dated 04-04-2019."),
        ],
        closing_lines=[
            "We humbly request you to please review the same and do the needful at your earliest convenience. This will enable us to remit the dues promptly and proceed with the Occupancy Certificate process.",
            "Thanking you,",
            "Yours respectfully,",
        ],
        signatory=[
            "Nishant Ranka",
            "Director,",
            "DRA Developers and Projects Pvt. Ltd.",
        ],
        output_path="/tmp/20260525_Ranka_Iris_License_Fee_Letter.docx"
    )
