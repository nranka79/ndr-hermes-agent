#!/opt/hermes/.venv/bin/python
"""
Beautified Market Research Report DOCX — DRAAS template.
Features: larger fonts, gold callout boxes, navy/gold branding, cover page,
radius-wise competitor tables, project profile sections, corridor comparison.

Usage: Copy this file, update paths/sections/data for your project, then run:
  /opt/hermes/.venv/bin/python /tmp/your_report.py

DRAAS colour palette:
  NAVY = #1B2A4A, GOLD = #C9A84C, GOLD_DARK = #A08430
  DARK_GRAY = #333333, MED_GRAY = #666666
  Callout BG = #FDF2D7, Table alt rows = #F5F7FA

Key sizing: body=11pt, table data=9pt, cover title=30pt, h1=navy, h2=gold-dark
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import datetime, os

# ── COLOURS ──
NAVY = RGBColor(0x1B, 0x2A, 0x4A)
GOLD = RGBColor(0xC9, 0xA8, 0x4C)
GOLD_DARK = RGBColor(0xA0, 0x84, 0x30)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
DG = RGBColor(0x33, 0x33, 0x33)      # dark grey body text
MG = RGBColor(0x66, 0x66, 0x66)      # medium grey for secondary
GRN = RGBColor(0x1E, 0x7A, 0x3C)      # green for positive
RED = RGBColor(0xC0, 0x39, 0x2B)      # red for warnings
AMB = RGBColor(0xE6, 0x7E, 0x22)      # amber for attention

# ── HELPERS ──
def shade(cell, hexc):
    cell._tc.get_or_add_tcPr().append(
        parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hexc}"/>'))

def cell_txt(cell, txt, bold=False, sz=10, color=DG, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = align
    r = p.add_run(str(txt))
    r.bold = bold; r.font.size = Pt(sz); r.font.color.rgb = color; r.font.name = 'Calibri'
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = Pt(13)

def table(doc, headers, rows, hdr_hex='1B2A4A', sz=9):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = True
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        shade(c, hdr_hex)
        cell_txt(c, h, bold=True, sz=sz+1, color=WHITE, align=WD_ALIGN_PARAGRAPH.CENTER)
    for ri, rd in enumerate(rows):
        for ci, ct2 in enumerate(rd):
            c = t.rows[ri+1].cells[ci]
            shade(c, 'F5F7FA' if ri%2==1 else 'FFFFFF')
            cell_txt(c, str(ct2), bold=(ci==0), sz=sz)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    return t

def h(doc, txt, lv=1):
    hd = doc.add_heading(txt, level=lv)
    for r in hd.runs:
        r.font.color.rgb = NAVY if lv==1 else (GOLD_DARK if lv==2 else NAVY)
        r.font.name = 'Calibri'
    return hd

def body(doc, txt, bold=False, sz=11, color=DG, sa=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(sa)
    r = p.add_run(txt)
    r.bold = bold; r.font.size = Pt(sz); r.font.color.rgb = color; r.font.name = 'Calibri'
    return p

def bullet(doc, txt, bp='', sz=10.5, ind=0.6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Cm(ind)
    if bp:
        r = p.add_run(bp); r.bold = True; r.font.size = Pt(sz); r.font.color.rgb = NAVY; r.font.name = 'Calibri'
    r = p.add_run(txt); r.font.size = Pt(sz); r.font.color.rgb = DG; r.font.name = 'Calibri'

def metric(doc, label, value):
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(3)
    r = p.add_run(f'{label}: '); r.bold = True; r.font.size = Pt(11)
    r.font.color.rgb = NAVY; r.font.name = 'Calibri'
    r = p.add_run(str(value)); r.font.size = Pt(11); r.font.color.rgb = DG; r.font.name = 'Calibri'

def callout(doc, txt):
    """Gold-background callout box for key insights."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Cm(0.5)
    r = p.add_run(f'▎ {txt}')
    r.font.size = Pt(11); r.font.color.rgb = DG; r.font.name = 'Calibri'; r.bold = True
    pPr = p._p.get_or_add_pPr()
    pPr.append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="FDF2D7" w:val="clear"/>'))

def hr(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('━'*60)
    r.font.color.rgb = RGBColor(0xDD, 0xDD, 0xDD)
    r.font.size = Pt(8)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)

def cover_page(doc, title, subtitle, location, date_str, doc_type):
    """Generate DRAAS-branded cover page."""
    for _ in range(4): doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('━'*60); r.font.color.rgb = GOLD; r.font.size = Pt(16)
    doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title); r.bold = True; r.font.size = Pt(30); r.font.color.rgb = NAVY; r.font.name = 'Calibri'
    doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(subtitle); r.font.size = Pt(18); r.font.color.rgb = GOLD_DARK; r.font.name = 'Calibri'
    doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(doc_type); r.font.size = Pt(14); r.font.color.rgb = MG; r.font.name = 'Calibri'
    doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('━'*60); r.font.color.rgb = GOLD; r.font.size = Pt(16)
    for _ in range(3): doc.add_paragraph()
    for lbl, val in [
        ('Prepared For', 'DRAAS — Real Estate & Infrastructure'),
        ('Location', location),
        ('Date', date_str),
        ('Type', doc_type),
    ]:
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(f'{lbl}:  '); r.bold = True; r.font.size = Pt(10); r.font.color.rgb = NAVY; r.font.name = 'Calibri'
        r = p.add_run(val); r.font.size = Pt(10); r.font.color.rgb = DG; r.font.name = 'Calibri'

def page_break(doc):
    doc.add_page_break()

def profile_table(doc, project_name, data_rows):
    """Create a 2-column key:value profile table for a project."""
    h(doc, project_name, 3)
    table(doc, ['Parameter', 'Details'], data_rows, sz=8.5)


# ── BUILD EXAMPLE ──
if __name__ == '__main__':
    doc = Document()
    for sec in doc.sections:
        sec.top_margin = Cm(2.5)
        sec.bottom_margin = Cm(2.0)
        sec.left_margin = Cm(2.8)
        sec.right_margin = Cm(2.8)
    
    sty = doc.styles['Normal']
    sty.font.name = 'Calibri'; sty.font.size = Pt(11); sty.font.color.rgb = DG

    cover_page(doc,
        title='MARKET RESEARCH &\nLOCATION ANALYSIS REPORT',
        subtitle='[Area] — [Corridor] | [District]',
        location='[Full address, coordinates]',
        date_str='[Month Year]',
        doc_type='Comprehensive Market Research & Location Analysis')

    page_break(doc)
    h(doc, '1. Executive Summary', 1)
    body(doc, 'Summary text here...', sz=12)
    hr(doc)
    for l, v in [('Key Metric 1', 'Value 1'), ('Key Metric 2', 'Value 2')]:
        metric(doc, l, v)
    hr(doc)

    page_break(doc)
    h(doc, '2. Location Identification & Site Analysis', 1)
    h(doc, '2.1 Site Details', 2)
    table(doc, ['Parameter', 'Details'], [
        ['Coordinates', '[lat, lon]'],
        ['Taluk', '[taluk]'],
        ['District', '[district]'],
    ])
    callout(doc, '★ KEY INSIGHT: Callout box for important findings.')

    # Continue with remaining sections...

    path = '/data/hermes/reports/Your_Report_Name.docx'
    doc.save(path)
    print(f'Saved: {path}')
    print(f'Size: {os.path.getsize(path)/1024:.1f} KB')
