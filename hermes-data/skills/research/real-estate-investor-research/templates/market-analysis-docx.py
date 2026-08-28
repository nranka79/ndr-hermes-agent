#!/opt/hermes/.venv/bin/python
"""
DOCX Market Analysis Report — COMPREHENSIVE TEMPLATE
=====================================================
Copy this file, then:
1. Update REPORT VARIABLES at the top
2. Fill in all tables with your research data
3. Run: /opt/hermes/.venv/bin/python /path/to/your_copy.py

Produces: DRAAS-branded DOCX with navy+gold, 18-section industry-standard structure.
Each section needs: narrative paragraph + supporting table + insight.
This is NOT a slides deck — write proper document prose.
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# ════════════════════════════════════════════
# REPORT VARIABLES — EDIT THESE
# ════════════════════════════════════════════
AREA = "Your Area"
CORRIDOR = "Key Corridor (e.g. NH44, Neighborhood Road)"
COORDS = "00°00'00.0\"N 00°00'00.0\"E"
DISTRICT = "District Name"
TALUK = "Taluk Name"
PLANNING_AUTHORITY = "City Corporation / LPA"
REPORT_DATE = "Month Year"
PLOT_PRICE_RANGE = "₹X,XXX–₹X,XXX/sqft"
LAUNCH_PRICE = "₹X,XXX–₹X,XXX/sqft"
UPSIDE = "XX–XX%"
ANCHOR_DEMAND = "X,000+ workforce within X km"
TOP_CATALYSTS = "Project 1 (timeline), Project 2 (timeline), Project 3 (timeline)"

# ── COLOURS (DRAAS Brand) ──
NAVY = RGBColor(0x1B, 0x2A, 0x4A)
NAVY_LIGHT = RGBColor(0x2C, 0x3E, 0x6B)
GOLD = RGBColor(0xC9, 0xA8, 0x4C)
GOLD_DARK = RGBColor(0xA0, 0x84, 0x30)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
BLACK = RGBColor(0x1D, 0x1F, 0x22)
DARK_GRAY = RGBColor(0x33, 0x33, 0x33)
MED_GRAY = RGBColor(0x66, 0x66, 0x66)
GREEN = RGBColor(0x1E, 0x7A, 0x3C)
RED = RGBColor(0xC0, 0x39, 0x2B)

# ════════════════════════════════════════════
# HELPER FUNCTIONS (copy verbatim)
# ════════════════════════════════════════════

def shade_cell(cell, color_hex):
    cell._tc.get_or_add_tcPr().append(
        parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>'))

def cell_text(cell, text, bold=False, size=8, color=DARK_GRAY, align=WD_ALIGN_PARAGRAPH.LEFT, font='Calibri'):
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = align
    r = p.add_run(str(text))
    r.bold = bold
    r.font.size = Pt(size)
    r.font.color.rgb = color
    r.font.name = font
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.line_spacing = Pt(10)

def make_table(doc, headers, rows, header_hex='1B2A4A', font_sz=7.5):
    """Create formatted table with navy header + alternating row colours."""
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = True
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        shade_cell(c, header_hex)
        cell_text(c, h, bold=True, size=font_sz+1, color=WHITE, align=WD_ALIGN_PARAGRAPH.CENTER)
    for ri, rd in enumerate(rows):
        for ci, ct in enumerate(rd):
            c = t.rows[ri+1].cells[ci]
            shade_cell(c, 'F8F9FA' if ri%2==1 else 'FFFFFF')
            cell_text(c, str(ct), bold=(ci==0), size=font_sz)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t

def section_h(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for r in h.runs:
        r.font.color.rgb = NAVY if level==1 else (GOLD_DARK if level==2 else NAVY)
        r.font.name = 'Calibri'
    return h

def body(doc, text, bold=False, size=9.5, color=DARK_GRAY, space_after=5):
    """Add a body paragraph with narrative text (not bullet)."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(size)
    r.font.color.rgb = color
    r.font.name = 'Calibri'
    return p

def bullet(doc, text, prefix='', size=9, indent=0.5):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.left_indent = Cm(indent)
    if prefix:
        r = p.add_run(prefix)
        r.bold = True; r.font.size = Pt(size); r.font.color.rgb = NAVY; r.font.name = 'Calibri'
    r = p.add_run(text)
    r.font.size = Pt(size); r.font.color.rgb = DARK_GRAY; r.font.name = 'Calibri'
    return p

def metric(doc, label, value):
    """Bold label: value — for key metric callouts."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(f'{label}: ')
    r.bold = True; r.font.size = Pt(10); r.font.color.rgb = NAVY; r.font.name = 'Calibri'
    r = p.add_run(value)
    r.font.size = Pt(10); r.font.color.rgb = DARK_GRAY; r.font.name = 'Calibri'

def hr(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('─' * 70)
    r.font.color.rgb = RGBColor(0xDD, 0xDD, 0xDD); r.font.size = Pt(6)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)

# ════════════════════════════════════════════
# BUILD DOCUMENT
# ════════════════════════════════════════════

doc = Document()
for sec in doc.sections:
    sec.top_margin = Cm(2.0)
    sec.bottom_margin = Cm(2.0)
    sec.left_margin = Cm(2.5)
    sec.right_margin = Cm(2.5)

style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(10)
style.font.color.rgb = DARK_GRAY

# ════════════════════════════════════════════
# COVER PAGE
# ════════════════════════════════════════════
for _ in range(5):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('━' * 55)
r.font.color.rgb = GOLD; r.font.size = Pt(14)

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('COMPREHENSIVE MARKET ANALYSIS &\nDEVELOPMENT FEASIBILITY REPORT')
r.bold = True; r.font.size = Pt(26); r.font.color.rgb = NAVY; r.font.name = 'Calibri'

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run(f'{AREA} — {CORRIDOR} | {DISTRICT} District')
r.font.size = Pt(16); r.font.color.rgb = GOLD_DARK; r.font.name = 'Calibri'

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Proposed Premium Plotted Layout Development\nwith Club House & Fully Loaded Amenities')
r.font.size = Pt(13); r.font.color.rgb = MED_GRAY; r.font.name = 'Calibri'

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('━' * 55)
r.font.color.rgb = GOLD; r.font.size = Pt(14)

for _ in range(3):
    doc.add_paragraph()

for lbl, val in [
    ('Prepared For', 'DRAAS — Real Estate & Infrastructure'),
    ('Subject Location', f'{AREA}, {CORRIDOR}, {DISTRICT} District, Karnataka'),
    ('Coordinates', COORDS),
    ('Report Date', REPORT_DATE),
    ('Document Type', 'Market Analysis & Development Feasibility (Industry Standard)'),
    ('Project Type', 'Premium Plotted Layout — Club House + Fully Loaded Amenities'),
]:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f'{lbl}:  ')
    r.bold = True; r.font.size = Pt(9); r.font.color.rgb = NAVY; r.font.name = 'Calibri'
    r = p.add_run(val)
    r.font.size = Pt(9); r.font.color.rgb = DARK_GRAY; r.font.name = 'Calibri'

doc.add_page_break()

# ════════════════════════════════════════════
# TABLE OF CONTENTS
# ════════════════════════════════════════════
section_h(doc, 'Table of Contents', 1)
toc = [
    '1. Executive Summary',
    '2. Location Identification & Site Analysis',
    '3. Regional Connectivity & Key Distances',
    '4. Demographic & Economic Profile',
    '5. Infrastructure Projects — Status & Timeline',
    '6. Employment Generators & Demand Drivers',
    '7. Development Regulations & Approvals',
    '8. Competitive Landscape — Nearby Projects (Radius-Wise)',
    '9. Competitive Pricing Analysis',
    '10. Social Infrastructure Mapping',
    '11. Demand Analysis & Target Buyer Persona',
    '12. Recommended Product Mix & Amenities Package',
    '13. Phasing Strategy & Pricing Roadmap',
    '14. Risk Assessment & Mitigation',
    '15. Investment Thesis & Exit Strategy',
    '16. Recommendation & Next Steps',
    '17. Sources & References',
    '18. Annexures',
]
for item in toc:
    p = doc.add_paragraph()
    r = p.add_run(item)
    r.font.size = Pt(10); r.font.color.rgb = NAVY; r.font.name = 'Calibri'
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Cm(1)

doc.add_page_break()

# ════════════════════════════════════════════
# 1. EXECUTIVE SUMMARY
# ════════════════════════════════════════════
section_h(doc, '1. Executive Summary', 1)

body(doc, (
    f'{AREA}, located on the {CORRIDOR}, represents a strategically positioned land parcel in '
    f'{DISTRICT} District. [Add 2-3 sentences on the opportunity — location, demand drivers, '
    f'market gap, and recommendation.]'
), size=10)

body(doc, 'This report provides a comprehensive assessment of the development potential, covering:')
bullets_exec = [
    'Location analysis and connectivity assessment across road, air, and rail networks',
    'Regulatory framework including planning authority jurisdiction, FAR, height restrictions, conversion requirements',
    'Detailed competitor mapping within 0–5 km, 5–10 km, and 10–20 km radii with actual project names and pricing',
    'Employment generators and demand driver analysis specific to the corridor',
    'Social infrastructure mapping (education, healthcare, retail, tourism)',
    'Development feasibility including recommended product mix, amenities, phasing, and pricing strategy',
    'Risk assessment covering regulatory, market, and execution risks',
]
for b in bullets_exec:
    bullet(doc, b, size=9.5)

hr(doc)

metric(doc, 'Subject Location', f'{AREA} ({CORRIDOR}), {DISTRICT} District')
metric(doc, 'Planning Authority', PLANNING_AUTHORITY)
metric(doc, 'Applicable Master Plan', '[BDA RMP 2031 / local master plan]')
metric(doc, 'Permissible FAR', '~X.XX–X.XX (residential plotted layout)')
metric(doc, 'Height Restriction', '[Within airport funnel — max G+X/G+X, or no restriction]')
metric(doc, 'Current Agricultural Land Rate', '₹XX–XX L/acre')
metric(doc, 'Current NA Converted Land Rate', '₹XX–XX L/acre')
metric(doc, 'Nearby Competing Plot Prices', PLOT_PRICE_RANGE)
metric(doc, 'Recommended Launch Price', LAUNCH_PRICE)
metric(doc, 'Projected Appreciation', UPSIDE)
metric(doc, 'Key Demand Anchors', ANCHOR_DEMAND)

doc.add_page_break()

# ════════════════════════════════════════════
# 2. LOCATION IDENTIFICATION & SITE ANALYSIS
# ════════════════════════════════════════════
section_h(doc, '2. Location Identification & Site Analysis', 1)

section_h(doc, '2.1 Site Coordinates & Administrative Hierarchy', 2)

make_table(doc,
    ['Parameter', 'Details'],
    [
        ['Coordinates', COORDS],
        ['Landmark', '[Nearest landmark / junction]'],
        ['Nearest Village', '[Village name]'],
        ['Taluk', TALUK],
        ['District', DISTRICT],
        ['State', 'Karnataka'],
        ['Pincode', '[Pincode]'],
        ['Sub-Registrar Office', '[Town] Sub-Registrar'],
        ['Planning Authority', PLANNING_AUTHORITY],
        ['Applicable Master Plan', '[Details]'],
        ['Airport Height Zone', '[Affected / Not affected]'],
    ],
    font_sz=8
)

section_h(doc, '2.2 Site Characteristics', 2)
body(doc, 'Describe the site: topography, frontage, current land use, surrounding context, visibility, scenic features.')

doc.add_page_break()

# ════════════════════════════════════════════
# 3. REGIONAL CONNECTIVITY
# ════════════════════════════════════════════
section_h(doc, '3. Regional Connectivity & Key Distances', 1)

section_h(doc, '3.1 Road Connectivity', 2)
body(doc, f'The subject location is well-connected via [primary corridors].')

make_table(doc,
    ['Route', 'Destination', 'Approx Distance', 'Typical Travel Time', 'Road Condition'],
    [
        ['[Highway] (South)', '[Destination 1]', '~X km', 'XX–XX min', '[Condition]'],
        ['[Highway] (North)', '[Destination 2]', '~X km', 'XX–XX min', '[Condition]'],
        ['[Local Road] (West)', '[Destination 3]', '~X km', 'XX–XX min', '[Condition]'],
    ],
    font_sz=7
)

section_h(doc, '3.2 Air Connectivity', 2)
body(doc, '[Describe airport proximity, airlines, expansion plans]')

section_h(doc, '3.3 Rail Connectivity', 2)
body(doc, '[Describe nearest station, suburban rail, metro status]')

section_h(doc, '3.4 Strategic Road Network', 2)
body(doc, '[Describe ring road, bypass, or expressway projects]')

doc.add_page_break()

# ════════════════════════════════════════════
# 4. DEMOGRAPHIC & ECONOMIC PROFILE
# ════════════════════════════════════════════
section_h(doc, '4. Demographic & Economic Profile', 1)

section_h(doc, '4.1 District Overview', 2)

make_table(doc,
    ['Parameter', 'Value'],
    [
        ['District Population', '[Census data]'],
        ['Projected Population', '[Estimate]'],
        ['Urban Population Share', '[%]'],
        ['Literacy Rate', '[%]'],
        ['Main Economic Activities', '[List]'],
        ['Key Towns', '[List]'],
    ],
    font_sz=8
)

section_h(doc, '4.2 Buyer Catchment Demographics', 2)
body(doc, 'The primary buyer catchment includes:')

catchments = [
    '[Employment Hub 1] — [count] workforce, [X] km',
    '[Employment Hub 2] — [count] workforce, [X] km',
    '[Employment Hub 3] — [count] workforce, [X] km',
]
for c in catchments:
    bullet(doc, c, '• ', size=9)

doc.add_page_break()

# ════════════════════════════════════════════
# 5. INFRASTRUCTURE PROJECTS
# ════════════════════════════════════════════
section_h(doc, '5. Infrastructure Projects — Status & Timeline', 1)

body(doc, 'The following infrastructure projects directly impact the subject micro-market.')

make_table(doc,
    ['Project', 'Category', 'Status', 'Timeline', 'Impact on Subject'],
    [
        ['[Project 1]', '[Road/Aviation/Rail]', '[Status]', '[Timeline]', '[Impact]'],
        ['[Project 2]', '[Category]', '[Status]', '[Timeline]', '[Impact]'],
        ['[Project 3]', '[Category]', '[Status]', '[Timeline]', '[Impact]'],
    ],
    font_sz=7
)

section_h(doc, '5.1 Infrastructure Impact Assessment', 2)
body(doc, '[Narrative analysis of how infrastructure projects will affect land values and demand]')

doc.add_page_break()

# ════════════════════════════════════════════
# 6. EMPLOYMENT GENERATORS & DEMAND DRIVERS
# ════════════════════════════════════════════
section_h(doc, '6. Employment Generators & Demand Drivers', 1)

section_h(doc, '6.1 Primary Employment Hubs (Within 20 km)', 2)

make_table(doc,
    ['Employment Hub', 'Distance', 'Sector', 'Key Employers', 'Est. Workforce', 'Demand Profile'],
    [
        ['[Hub 1]', '~X km', '[Sector]', '[Employers]', '[Count]', '[Profile]'],
        ['[Hub 2]', '~X km', '[Sector]', '[Employers]', '[Count]', '[Profile]'],
    ],
    font_sz=7
)

section_h(doc, '6.2 Secondary Employment Hubs (Within 25–35 km)', 2)
body(doc, '[Table or list of secondary hubs]')

section_h(doc, '6.3 Demand Driver Analysis', 2)

demand_drivers = [
    ('[Driver 1 Title]', '[2-3 sentence analysis of why this driver matters for the proposed development]'),
    ('[Driver 2 Title]', '[Analysis]'),
    ('[Driver 3 Title]', '[Analysis]'),
]
for title, desc in demand_drivers:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(1)
    r = p.add_run(title)
    r.bold = True; r.font.size = Pt(9); r.font.color.rgb = NAVY; r.font.name = 'Calibri'
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.left_indent = Cm(0.5)
    r = p.add_run(desc)
    r.font.size = Pt(9); r.font.color.rgb = DARK_GRAY; r.font.name = 'Calibri'

doc.add_page_break()

# ════════════════════════════════════════════
# 7. DEVELOPMENT REGULATIONS & APPROVALS
# ════════════════════════════════════════════
section_h(doc, '7. Development Regulations & Approvals', 1)

body(doc, (
    'This section is CRITICAL as it establishes the legal and regulatory framework. '
    'Errors in understanding the regulatory context can lead to plan rejection, project delays, or legal challenges.'
), size=9.5, bold=True, color=RED)

section_h(doc, '7.1 Planning Authority & Jurisdiction', 2)

make_table(doc,
    ['Authority', 'Role', 'Relevance'],
    [
        ['[Primary LPA]', 'Building plan and layout approval', 'Primary sanctioning authority'],
        ['[DC Office]', 'Revenue conversion under KLR Act', 'Agri to NA conversion'],
        ['[Master Plan Authority]', 'Land-use zoning', 'Determines permissible use'],
        ['[AAI / DGCA]', 'Airport height restriction', '[Affected / Not affected]'],
    ],
    font_sz=7
)

section_h(doc, '7.2 Land Use & Master Plan Zoning', 2)
body(doc, '[Describe BDA RMP 2031 or equivalent zoning]')

section_h(doc, '7.3 FAR & Building Regulations', 2)

make_table(doc,
    ['Parameter', 'Residential Plotted Layout', 'Commercial / Mixed-Use'],
    [
        ['Maximum FAR', '~X.XX–X.XX', '~X.XX–X.XX'],
        ['Ground Coverage', '~XX–XX%', '~XX–XX%'],
        ['Maximum Height', '~X m (G+X/G+X)', '~X m'],
        ['Minimum Plot Size', 'XXX sqm per KTCP Act', 'As per LPA norms'],
        ['Minimum Road Width', 'X–X m', 'X m+'],
        ['Open Space Requirement', 'XX–XX% of layout area', 'As per LPA norms'],
    ],
    font_sz=7
)

section_h(doc, '7.4 Conversion Process (Agri to NA)', 2)
body(doc, 'Sequential approval process:')
conversion_steps = [
    'Step 1: CLU (Change of Land Use) — apply to [authority]',
    'Step 2: Revenue Conversion — apply to DC under Section 95 of KLR Act',
    'Step 3: Layout Plan Approval — submit to [LPA]',
    'Step 4: Building Plan Approval — submit for each structure',
    'Step 5: RERA Registration — mandatory before marketing (>8 plots)',
]
for s in conversion_steps:
    bullet(doc, s, '• ', size=9)

section_h(doc, '7.5 Critical Height Restriction', 2)
body(doc, '[Describe airport approach funnel restriction if applicable, or note "not affected"]', bold=True)

doc.add_page_break()

# ════════════════════════════════════════════
# 8. COMPETITIVE LANDSCAPE (RADIUS-WISE)
# ════════════════════════════════════════════
section_h(doc, '8. Competitive Landscape — Nearby Projects (Radius-Wise)', 1)

# Zone 1: 0-5km
section_h(doc, '8.1 Zone 1: Immediate Vicinity (0–5 km Radius)', 2)
body(doc, '[Describe what exists in this zone — note if zero direct competition]')

make_table(doc,
    ['Project', 'Developer', 'Type', 'Plot Sizes (sqft)', 'Price (₹/sqft)', 'Status', 'Distance'],
    [
        ['[Project]', '[Developer]', '[Type]', '[Sizes]', 'X,XXX–X,XXX', '[Status]', '~X km'],
    ],
    font_sz=7
)

# Zone 2: 5-10km
section_h(doc, '8.2 Zone 2: Primary Competitive Zone (5–10 km)', 2)
body(doc, 'This is the PRIMARY COMPETITIVE ZONE where most comparable projects exist.')

make_table(doc,
    ['Project', 'Developer', 'Type', 'Plot Sizes (sqft)', 'Price (₹/sqft)', 'Club House', 'Pool', 'Status', 'Distance'],
    [
        ['[Project]', '[Developer]', '[Type]', '[Sizes]', 'X,XXX–X,XXX', '✓/-', '✓/-', '[Status]', '~X km'],
        ['[Project]', '[Developer]', '[Type]', '[Sizes]', 'X,XXX–X,XXX', '✓/-', '✓/-', '[Status]', '~X km'],
    ],
    font_sz=7
)

# Zone 3: 10-20km
section_h(doc, '8.3 Zone 3: Extended Competitive Radius (10–20 km)', 2)
body(doc, 'Indirect competition — different micro-market. Include for corridor price benchmarking only.')

make_table(doc,
    ['Project', 'Developer', 'Type', 'Price (₹/sqft)', 'Status', 'Distance'],
    [
        ['[Project]', '[Developer]', '[Type]', 'X,XXX–X,XXX', '[Status]', '~X km'],
    ],
    font_sz=7
)

section_h(doc, '8.4 Competitive Positioning Summary', 2)
body(doc, '[Synthesis — what is the gap / white space / first-mover opportunity]')

doc.add_page_break()

# ════════════════════════════════════════════
# 9. COMPETITIVE PRICING ANALYSIS
# ════════════════════════════════════════════
section_h(doc, '9. Competitive Pricing Analysis', 1)

section_h(doc, '9.1 Price Benchmarking by Corridor', 2)

make_table(doc,
    ['Corridor / Micro-Market', 'Plot Price Range (₹/sqft)', 'YoY Appreciation', 'Distance', 'Maturity Stage'],
    [
        ['[Corridor 1]', 'X,XXX–X,XXX', '+XX–XX%', 'X–X km', '[Stage]'],
        ['[Subject Corridor]', 'X,XXX–X,XXX', '+XX–XX%', 'Subject', '[Stage]'],
        ['[Corridor 2]', 'X,XXX–X,XXX', '+XX–XX%', 'X–X km', '[Stage]'],
    ],
    font_sz=7
)

section_h(doc, '9.2 Price Segmentation', 2)

make_table(doc,
    ['Segment', '₹/sqft', 'Comparative Projects', 'Target Buyer Profile'],
    [
        ['Entry-level', 'X,XXX–X,XXX', '[Projects]', '[Profile]'],
        ['Mid-market (gated, standard amenities)', 'X,XXX–X,XXX', '[Projects]', '[Profile]'],
        ['Premium (club house, pool, gym)', 'X,XXX–X,XXX', '[Projects — or note gap]', '[Profile]'],
        ['Ultra-premium', 'X,XXX–X,XXX+', '[Projects]', '[Profile]'],
    ],
    font_sz=7
)

section_h(doc, '9.3 White Space & Opportunity Gap', 2)
body(doc, '[Describe the gap the proposed project will fill]')

doc.add_page_break()

# ════════════════════════════════════════════
# 10. SOCIAL INFRASTRUCTURE
# ════════════════════════════════════════════
section_h(doc, '10. Social Infrastructure Mapping', 1)

section_h(doc, '10.1 Schools', 2)
make_table(doc,
    ['School', 'Board', 'Distance', 'Notes'],
    [
        ['[School]', 'CBSE/ICSE', '~X km', '[Notes]'],
    ],
    font_sz=7
)

section_h(doc, '10.2 Hospitals', 2)
make_table(doc,
    ['Hospital', 'Type', 'Distance', 'Key Features'],
    [
        ['[Hospital]', '[Type]', '~X km', '[Features]'],
    ],
    font_sz=7
)

section_h(doc, '10.3 Shopping & Entertainment', 2)
make_table(doc,
    ['Destination', 'Type', 'Distance', 'Details'],
    [
        ['[Place]', '[Type]', '~X km', '[Details]'],
    ],
    font_sz=7
)

section_h(doc, '10.4 Tourism & Lifestyle', 2)
make_table(doc,
    ['Attraction', 'Type', 'Distance', 'Selling Point'],
    [
        ['[Attraction]', '[Type]', '~X km', '[Selling point]'],
    ],
    font_sz=7
)

doc.add_page_break()

# ════════════════════════════════════════════
# 11. DEMAND ANALYSIS & TARGET BUYER
# ════════════════════════════════════════════
section_h(doc, '11. Demand Analysis & Target Buyer Persona', 1)

section_h(doc, '11.1 Buyer Segment Analysis', 2)

make_table(doc,
    ['Buyer Segment', 'Income Profile', 'Budget (₹)', 'Decision Driver', 'Demand Share'],
    [
        ['[Segment 1]', '₹XX–XX L/yr', 'XX–XX L', '[Driver]', 'XX%'],
        ['[Segment 2]', '₹XX–XX L/yr', 'XX–XX L', '[Driver]', 'XX%'],
    ],
    font_sz=7
)

section_h(doc, '11.2 Demand Volume Estimate', 2)

make_table(doc,
    ['Source', 'Workforce (est.)', '% Likely to Buy', 'Annual Housing Demand'],
    [
        ['[Source 1]', '[Count]', 'X%', '~X units/yr'],
        ['[Source 2]', '[Count]', 'X%', '~X units/yr'],
        ['Total', '', '', '~X,XXX units/yr'],
    ],
    font_sz=7
)

doc.add_page_break()

# ════════════════════════════════════════════
# 12. PRODUCT MIX & AMENITIES
# ════════════════════════════════════════════
section_h(doc, '12. Recommended Product Mix & Amenities Package', 1)

section_h(doc, '12.1 Proposed Layout Configuration', 2)
body(doc, 'Assuming a developable area of X acres:')

make_table(doc,
    ['Plot Size (sqft)', 'Count', '%', 'Target Buyer', 'Est. Price (₹/sqft)', 'Est. Total Value (₹ Cr)'],
    [
        ['X,XXX–X,XXX', 'XX', 'XX%', '[Profile]', 'X,XXX–X,XXX', '~XX–XX'],
        ['X,XXX–X,XXX', 'XX', 'XX%', '[Profile]', 'X,XXX–X,XXX', '~XX–XX'],
        ['Total', 'XXX', '100%', '', '', '~XX–XX Cr (GDV)'],
    ],
    font_sz=7
)

section_h(doc, '12.2 Amenities Package', 2)

make_table(doc,
    ['Amenity', 'Specification', 'Est. Cost (₹ Cr)', 'Marketing Impact'],
    [
        ['Club House', 'X,XXX–X,XXX sqft — [description]', '~X.X–X.X', '[Impact]'],
        ['Swimming Pool', '[Spec]', '~X.X–X.X', '[Impact]'],
        ['Gym', '[Spec]', '~X.X–X.X', '[Impact]'],
        ['Central Park', '[Spec]', '~X.X–X.X', '[Impact]'],
        ['Security Infrastructure', '[Spec]', '~X.X–X.X', '[Impact]'],
        ['Total', '', '~X.X–X.X Cr', ''],
    ],
    font_sz=7
)

doc.add_page_break()

# ════════════════════════════════════════════
# 13. PHASING STRATEGY & PRICING ROADMAP
# ════════════════════════════════════════════
section_h(doc, '13. Phasing Strategy & Pricing Roadmap', 1)

section_h(doc, '13.1 Development Phasing', 2)

make_table(doc,
    ['Phase', 'Timeline', 'Plots Released', 'Key Activities', 'Investment Trigger'],
    [
        ['Phase 0: Pre-Launch', 'Month 1–3', 'XX', '[Activities]', 'Early bird bookings'],
        ['Phase 1: Launch', 'Month 4–6', 'XX', '[Activities]', 'Main launch pricing'],
        ['Phase 2: Growth', 'Month 7–12', 'XX', '[Activities]', 'Amenities premium'],
        ['Phase 3: Harvest', 'Year 2', 'XX', '[Activities]', 'Ready-to-build premium'],
    ],
    font_sz=7
)

section_h(doc, '13.2 Pricing Roadmap', 2)

make_table(doc,
    ['Milestone', '₹/sqft', 'Cumulative Escalation', 'Rationale'],
    [
        ['Early Bird', 'X,XXX', 'Base', '[Rationale]'],
        ['Phase 1 Launch', 'X,XXX', '+XX%', '[Rationale]'],
        ['Phase 2', 'X,XXX', '+XX%', '[Rationale]'],
        ['Phase 3 Ready-to-Build', 'X,XXX', '+XX%', '[Rationale]'],
    ],
    font_sz=7
)

section_h(doc, '13.3 Indicative Development Budget', 2)

make_table(doc,
    ['Cost Head', 'Total (₹ Cr)', '% of GDV'],
    [
        ['Land Acquisition (X acres @ ₹X L/acre)', 'X.XX', 'X–X%'],
        ['Land Conversion & Approvals', 'X.XX–X.XX', 'X%'],
        ['Site Development', 'X.X–X.X', 'X%'],
        ['Amenities', 'X.X–X.X', 'X%'],
        ['Marketing, Sales & Brokerage', 'X.X–X.X', 'X%'],
        ['Finance Cost', 'X.X–X.X', 'X%'],
        ['Contingency', 'X.XX–X.XX', 'X–X%'],
        ['Total Development Cost', '~XX–XX Cr', 'XX–XX% of GDV'],
    ],
    font_sz=7
)

doc.add_page_break()

# ════════════════════════════════════════════
# 14. RISK ASSESSMENT
# ════════════════════════════════════════════
section_h(doc, '14. Risk Assessment & Mitigation', 1)

make_table(doc,
    ['Risk Category', 'Risk Description', 'Severity', 'Probability', 'Mitigation Strategy'],
    [
        ['Regulatory', '[Risk]', 'High/Med/Low', 'High/Med/Low', '[Mitigation]'],
        ['Market', '[Risk]', 'High/Med/Low', 'High/Med/Low', '[Mitigation]'],
        ['Infrastructure', '[Risk]', 'High/Med/Low', 'High/Med/Low', '[Mitigation]'],
        ['Execution', '[Risk]', 'High/Med/Low', 'High/Med/Low', '[Mitigation]'],
        ['Financial', '[Risk]', 'High/Med/Low', 'High/Med/Low', '[Mitigation]'],
    ],
    font_sz=7
)

body(doc, '[Overall risk rating and summary]')

doc.add_page_break()

# ════════════════════════════════════════════
# 15. INVESTMENT THESIS
# ════════════════════════════════════════════
section_h(doc, '15. Investment Thesis & Exit Strategy', 1)

section_h(doc, '15.1 Investment Thesis Summary', 2)

make_table(doc,
    ['Parameter', 'Base Case', 'Bull Case', 'Bear Case'],
    [
        ['Land Acquisition Cost', '₹X L/acre', '₹X L/acre', '₹X L/acre'],
        ['Average Sale Price', '₹X,XXX/sqft', '₹X,XXX/sqft', '₹X,XXX/sqft'],
        ['Sales Period', 'XX months', 'XX months', 'XX months'],
        ['Development Cost', '₹X Cr', '₹X Cr', '₹X Cr'],
        ['Project GDV', '~₹X Cr', '~₹X Cr', '~₹X Cr'],
        ['Gross Margin', '~XX%', '~XX%', '~XX%'],
    ],
    font_sz=7
)

section_h(doc, '15.2 Exit Scenarios', 2)

make_table(doc,
    ['Exit Route', 'Timeline', 'Expected Return', 'Buyer Type', 'Feasibility'],
    [
        ['Retail sell-through (individual plots)', 'XX–XX months', 'XX–XX% ROI', 'End-users + investors', 'Primary route'],
        ['Bulk sale to developer', 'Month X–X', 'XX–XX% ROI', 'Institutional buyer', 'Feasible'],
        ['Hybrid: retail + bulk', 'XX–XX months', 'XX–XX% ROI', 'Mixed', 'Optimal'],
        ['Land flip (pre-approvals)', 'Month X–X', 'XX–XX% ROI', 'Developer', 'Min effort'],
    ],
    font_sz=7
)

section_h(doc, '15.3 Corridor Comparison', 2)

make_table(doc,
    ['Corridor', 'Plot Price (₹/sqft)', 'Entry Yield', '18-month Outlook', 'Risk-Reward'],
    [
        ['Subject', 'X,XXX–X,XXX', '~XX–XX%', '+XX–XX%', '★★★★★'],
        ['[Corridor 2]', 'X,XXX–X,XXX', '~XX–XX%', '+XX–XX%', '★★★★'],
        ['[Corridor 3]', 'X,XXX–X,XXX', '~XX–XX%', '+XX–XX%', '★★★'],
    ],
    font_sz=7
)

doc.add_page_break()

# ════════════════════════════════════════════
# 16. RECOMMENDATION & NEXT STEPS
# ════════════════════════════════════════════
section_h(doc, '16. Recommendation & Next Steps', 1)

body(doc, (
    'Based on this comprehensive analysis, the subject location presents a [STRONG / MODERATE / WEAK] '
    'opportunity for a premium plotted layout development.'
), size=10, bold=True)

section_h(doc, '16.1 Immediate Next Steps', 2)

next_steps = [
    'Survey number verification at Sub-Registrar office for encumbrances and title.',
    'Master plan zoning confirmation from the planning authority.',
    'AAI height NOC pre-consultation (if applicable).',
    'Land conversion feasibility assessment with local consultant.',
    'Topographical survey and soil testing.',
    'Preliminary layout plan preparation with registered architect.',
    'Competitor site visits to assess actual quality and sales velocity.',
    'First-party demand survey with target employer base.',
]
for i, step in enumerate(next_steps, 1):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Cm(0.5)
    r = p.add_run(f'{i}. {step}')
    r.font.size = Pt(9); r.font.color.rgb = DARK_GRAY; r.font.name = 'Calibri'

section_h(doc, '16.2 Strategic Recommendation', 2)

recs = [
    ('Phase 1 — Land Assemblage', '[Description]'),
    ('Phase 2 — Approvals', '[Description]'),
    ('Phase 3 — Launch', '[Description]'),
    ('Phase 4 — Scale', '[Description]'),
]
for title, desc in recs:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(1)
    r = p.add_run(title)
    r.bold = True; r.font.size = Pt(9.5); r.font.color.rgb = NAVY; r.font.name = 'Calibri'
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.left_indent = Cm(0.5)
    r = p.add_run(desc)
    r.font.size = Pt(9); r.font.color.rgb = DARK_GRAY; r.font.name = 'Calibri'

doc.add_paragraph()
hr(doc)

# ════════════════════════════════════════════
# 17. SOURCES & REFERENCES
# ════════════════════════════════════════════
section_h(doc, '17. Sources & References', 1)

sources = [
    '[Source 1] — [description]',
    '[Source 2] — [description]',
]
for s in sources:
    p = doc.add_paragraph()
    r = p.add_run(f'• {s}')
    r.font.size = Pt(9); r.font.color.rgb = DARK_GRAY; r.font.name = 'Calibri'
    p.paragraph_format.space_after = Pt(2)

# ════════════════════════════════════════════
# FOOTER
# ════════════════════════════════════════════
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('END OF REPORT')
r.bold = True; r.font.size = Pt(11); r.font.color.rgb = NAVY; r.font.name = 'Calibri'
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('DRAAS — Real Estate & Infrastructure  |  Bangalore  |  Chennai')
r.font.size = Pt(9); r.font.color.rgb = MED_GRAY; r.font.name = 'Calibri'
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Disclaimer: This report is prepared for DRAAS internal use. Competitor data is based on publicly available information and market estimates. Actual pricing, availability, and regulatory status should be independently verified.')
r.font.size = Pt(7); r.font.color.rgb = MED_GRAY; r.font.name = 'Calibri'
r.italic = True

# ════════════════════════════════════════════
# SAVE
# ════════════════════════════════════════════
output = f'/data/hermes/reports/{AREA.replace(\" \", \"_\")}_Market_Analysis_Comprehensive.docx'
doc.save(output)
import os
print(f'DOCX saved: {output}')
print(f'Size: {os.path.getsize(output) / 1024:.1f} KB')
